from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
from typing import List, Dict, Optional
from .rule_engine import RuleEngine

ABSTRACT_CLEAN_RE = re.compile(r'^[\t \u2022\u00b7\u2027\-\u3000]*(.*?)[\t ]*$')
SKIP_SECTIONS = {'摘要', 'abstract', 'abSTRACT', '目錄', 'table of contents',
                 '參考文獻', 'references', '誌謝', 'acknowledgements',
                 '附錄', 'appendix', '自傳', '符號說明',
                 '圖目錄', '表目錄', 'list of figures', 'list of tables'}
NUMBERED_LIST_RE = re.compile(r'^[\t ]*(\d{1,2})[.、）)]\s')

class FormatAnalyzer:
    def __init__(self, rule_file: str = "thesis_zh.yaml"):
        self.rule_engine = RuleEngine(rule_file)
        self.issues = []
    
    def analyze(self, doc_path: str) -> List[Dict]:
        doc = Document(doc_path)
        self.issues = []
        
        self._check_page_margins(doc)
        self._check_heading_page_break(doc)
        self._check_line_spacing(doc)
        self._check_numbering_xml(doc)
        self._check_first_line_indent(doc)
        self._check_page_numbers(doc)
        self._check_caption_format(doc)
        self._check_toc_styles(doc)
        self._check_empty_paragraphs(doc)
        self._check_direct_font(doc)
        self._check_abstract(doc)
        self._check_keywords(doc)
        
        return self.issues
    
    def _clean(self, text: str) -> str:
        m = ABSTRACT_CLEAN_RE.match(text)
        return m.group(1).strip() if m else text.strip()
    
    def _check_page_margins(self, doc):
        margins = self.rule_engine.get_page_margins("body")
        for i, section in enumerate(doc.sections):
            for label, actual, expected in [
                ('上邊距', section.top_margin.cm, margins[0]),
                ('下邊距', section.bottom_margin.cm, margins[1]),
                ('左邊距', section.left_margin.cm, margins[2]),
                ('右邊距', section.right_margin.cm, margins[3]),
            ]:
                if abs(actual - expected) > 0.2:
                    self.issues.append({
                        'type': 'page_margin', 'severity': 'error',
                        'message': f'第{i+1}節{label}: {actual:.1f}cm (應為{expected:.1f}cm)',
                        'fixable': True
                    })
    
    def _check_heading_page_break(self, doc):
        h1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1' and p.text.strip()]
        
        for para in h1_paras:
            pPr = para._element.find(qn('w:pPr'))
            has_break = False
            
            if pPr is not None:
                pb = pPr.find(qn('w:pageBreakBefore'))
                if pb is not None and pb.get(qn('w:val')) == '1':
                    has_break = True
            
            if not has_break:
                self.issues.append({
                    'type': 'heading_page_break', 'severity': 'warning',
                    'message': f'章節標題前應插入分頁符號: "{para.text.strip()[:25]}"',
                    'fixable': True
                })
    
    def _check_line_spacing(self, doc):
        expected = self.rule_engine.get('line_spacing.value', 1.5)
        for i, para in enumerate(doc.paragraphs):
            if para._element.getparent().tag.endswith('}tc'):
                continue
            pf = para.paragraph_format
            rule = pf.line_spacing_rule
            ls = pf.line_spacing
            
            if (
                rule == WD_LINE_SPACING.ONE_POINT_FIVE  # 行距恰為1.5倍（python-docx對360twips的簡化）
                or (
                    rule == WD_LINE_SPACING.MULTIPLE
                    and isinstance(ls, (int, float))
                    and abs(ls - expected) < 0.1
                )
            ):
                continue
            self.issues.append({
                'type': 'line_spacing', 'severity': 'warning',
                'message': f'段落{i+1}行距不符 (應為{expected}倍)',
                'fixable': True
            })
    
    def _check_numbering_xml(self, doc):
        numbering_part = doc.part.numbering_part
        if not numbering_part:
            return
        
        dash_count = 0
        non_tnr = 0
        
        for abs_num in numbering_part._element.findall(qn('w:abstractNum')):
            for lvl in abs_num.findall(qn('w:lvl')):
                txt_el = lvl.find(qn('w:lvlText'))
                if txt_el is not None:
                    val = txt_el.get(qn('w:val'))
                    if val and '-' in val:
                        dash_count += 1
                
                if txt_el is not None and txt_el.get(qn('w:val')):
                    rPr = lvl.find(qn('w:rPr'))
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            if rFonts.get(qn('w:ascii')) != 'Times New Roman':
                                non_tnr += 1
                        else:
                            non_tnr += 1
                    else:
                        non_tnr += 1
        
        if dash_count:
            self.issues.append({
                'type': 'numbering_xml', 'severity': 'warning',
                'message': f'編號使用"-"格式: {dash_count}處',
                'fixable': True
            })
        if non_tnr:
            self.issues.append({
                'type': 'numbering_font', 'severity': 'warning',
                'message': f'編號數字非 Times New Roman: {non_tnr}處',
                'fixable': True
            })
    
    def _check_first_line_indent(self, doc):
        cfg = self.rule_engine.get('first_line_indent', {})
        if not cfg.get('enabled', True):
            return
        expected_cm = cfg.get('value_cm', 0.75)
        skip_styles = {
            'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title',
            'table of figures', 'table of tables', 'toc',
        }
        in_cover = True
        count_bad = 0
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt in ('摘要', 'Abstract', '致謝', '目錄', '圖目錄', '表目錄', '參考文獻', '目次'):
                in_cover = False
            if in_cover:
                continue
            if para._element.getparent().tag.endswith('}tc'):
                continue
            sname = para.style.name if para.style else ''
            if sname in skip_styles:
                continue
            if re.match(r'^(?:[图表圖表]|\b[Ff]ig(?:ure)?\.?)\s*\d+[-.]\d+', txt):
                continue
            if not txt:
                continue
            pf = para.paragraph_format
            actual = pf.first_line_indent
            if actual is None:
                count_bad += 1
            elif abs(actual.cm - expected_cm) > 0.1:
                count_bad += 1
        if count_bad:
            self.issues.append({
                'type': 'first_line_indent', 'severity': 'warning',
                'message': f'{count_bad}段缺少首行縮排或值不正確（應為{expected_cm}cm）',
                'fixable': True
            })

    def _check_page_numbers(self, doc):
        for i, section in enumerate(doc.sections):
            footer = section.footer
            if not footer.paragraphs:
                self.issues.append({
                    'type': 'page_number', 'severity': 'warning',
                    'message': f'第{i+1}節 footer 為空，無頁碼',
                    'fixable': True
                })
                continue
            para = footer.paragraphs[0]
            has_page = False
            for r in para._element.findall(qn('w:r')):
                if r.find(qn('w:fldChar')) is not None or r.find(qn('w:instrText')) is not None:
                    has_page = True
                    break
            if not has_page:
                self.issues.append({
                    'type': 'page_number', 'severity': 'warning',
                    'message': f'第{i+1}節未發現頁碼',
                    'fixable': True
                })

    def _check_caption_format(self, doc):
        CAPTION_RE = re.compile(r'((?:[图表圖表]|\b[Ff]ig(?:ure)?\.?))\s*(\d+)-(\d+)')
        count = 0
        for para in doc.paragraphs:
            if CAPTION_RE.search(para.text):
                count += 1
        if count:
            self.issues.append({
                'type': 'caption_format', 'severity': 'warning',
                'message': f'{count}處圖表編號使用 "-"（應為 "."）',
                'fixable': True
            })

    def _check_toc_styles(self, doc):
        toc_styles = 0
        has_hyperlink = False
        for s in doc.styles:
            nl = s.name.lower()
            if nl in ('toc 1', 'toc 2', 'toc 3', 'table of figures'):
                toc_styles += 1
                rPr = s._element.find(qn('w:rPr'))
                if rPr is not None:
                    rStyle = rPr.find(qn('w:rStyle'))
                    if rStyle is not None and rStyle.get(qn('w:val'), '').lower() == 'hyperlink':
                        has_hyperlink = True
                    u = rPr.find(qn('w:u'))
                    if u is not None:
                        has_hyperlink = True
        if not toc_styles:
            self.issues.append({
                'type': 'toc_style', 'severity': 'info',
                'message': '未發現目錄樣式（如無目錄可忽略）',
                'fixable': True
            })
        elif has_hyperlink:
            self.issues.append({
                'type': 'toc_style', 'severity': 'warning',
                'message': '目錄樣式包含超連結格式（底線/藍色/字元樣式）',
                'fixable': True
            })

    def _check_empty_paragraphs(self, doc):
        count = 0
        for para in doc.paragraphs:
            if para._element.getparent().tag.endswith('}tc'):
                continue
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and 'Heading' in pStyle.get(qn('w:val'), ''):
                    continue
            text_content = ''.join(t.text or '' for t in para._element.findall('.//' + qn('w:t')))
            if not text_content.strip():
                if para._element.findall('.//' + qn('w:drawing')):
                    continue
                count += 1
        if count > 5:
            self.issues.append({
                'type': 'empty_paragraphs', 'severity': 'warning',
                'message': f'發現{count}個空段落（多餘空白行）',
                'fixable': True
            })

    def _check_direct_font(self, doc):
        count = 0
        for para in doc.paragraphs:
            for run in para._element.findall(qn('w:r')):
                rPr = run.find(qn('w:rPr'))
                if rPr is None:
                    continue
                if rPr.find(qn('w:rFonts')) is not None:
                    count += 1
                    break
        if count:
            self.issues.append({
                'type': 'direct_font', 'severity': 'info',
                'message': f'{count}段落包含直接字型設定（樣式可能無法生效）',
                'fixable': True
            })

    def _check_abstract(self, doc):
        found = {'摘要': False, 'abstract': False}
        for para in doc.paragraphs:
            clean = self._clean(para.text)
            if clean in found:
                found[clean] = True
            elif clean.upper() == 'ABSTRACT':
                found['abstract'] = True
        
        if not found['摘要']:
            self.issues.append({
                'type': 'missing_section', 'severity': 'warning',
                'message': '未找到"摘要"章節標題', 'fixable': False
            })
        if not found['abstract']:
            self.issues.append({
                'type': 'missing_section', 'severity': 'warning',
                'message': '未找到"Abstract"章節標題', 'fixable': False
            })
    
    def _check_keywords(self, doc):
        KW_RE = re.compile(r'^(關鍵詞|關鍵字|Key words|Keywords)\s*[:：]?\s*', re.IGNORECASE)
        found = False
        for para in doc.paragraphs:
            if KW_RE.match(para.text.strip()):
                found = True
                break
        if not found:
            self.issues.append({
                'type': 'missing_keywords', 'severity': 'warning',
                'message': '未找到"關鍵詞/Keywords"段落', 'fixable': True
            })
