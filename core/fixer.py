from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import traceback
from pathlib import Path
from typing import List, Dict
from .rule_engine import RuleEngine
from .debug_log import get_logger

log = get_logger("fixer")

class FormatFixer:
    def __init__(self, rule_file: str = "thesis_zh.yaml"):
        self.rule_engine = RuleEngine(rule_file)
    
    def fix_document(self, input_path: str, output_path: str) -> List[Dict]:
        log.info("=== fix_document 開始 ===")
        log.info("input_path=%s, output_path=%s", input_path, output_path)
        
        if not Path(input_path).exists():
            msg = f"輸入檔案不存在: {input_path}"
            log.error(msg)
            return [{'type': 'error', 'message': msg}]
        
        doc = Document(input_path)
        fixes = []
        
        fix_methods = [
            ('page_margins', self._fix_page_margins),
            ('heading_style_defs', self._fix_heading_style_defs),
            ('direct_font_formatting', self._fix_direct_font_formatting),
            ('heading_indent', self._remove_heading_indent),
            ('heading_page_break', self._fix_heading_page_break),
            ('special_sections_page_break', self._fix_special_sections_page_break),
            ('line_spacing', self._fix_line_spacing),
            ('first_line_indent', self._fix_first_line_indent),
            ('keywords_format', self._fix_keywords_format),
            ('numbering_xml', self._fix_numbering_xml),
            ('empty_paragraphs', self._fix_empty_paragraphs),
            ('toc_hyperlinks', self._fix_toc_hyperlinks),
            ('caption_numbering', self._fix_caption_numbering),
            ('heading_numbering', self._fix_heading_numbering),
            ('caption_alignment', self._fix_caption_alignment),
            ('table_alignment', self._fix_table_alignment),
            ('table_no_break_across_pages', self._fix_table_no_break_across_pages),
            ('toc_field_codes', self._fix_toc_field_codes),
            ('references_numbered_list', self._convert_references_to_numbered_list),
            ('cross_references', self._fix_cross_references),
            ('inline_citations_xref', self._convert_inline_citations_to_xrefs),
            ('toc_style', self._fix_toc_style),
            ('cover_tabs', self._fix_cover_tabs),
            ('page_numbers', self._fix_page_numbers),
        ]
        
        for name, method in fix_methods:
            try:
                log.info(">>> 開始執行: %s", name)
                result = method(doc)
                fixes.extend(result)
                log.info("<<< %s 完成, 回傳 %d 條結果", name, len(result))
            except Exception as e:
                log.error("!!! %s 發生例外: %s", name, e)
                log.error(traceback.format_exc())
                fixes.append({'type': name, 'message': f'執行錯誤: {e}'})
        
        log.info("所有 fix 執行完畢, 共 %d 條結果", len(fixes))
        log.info("正在儲存檔案: %s", output_path)
        output = Path(output_path)
        doc.save(str(output))
        log.info("檔案已儲存, 大小=%d bytes", output.stat().st_size if output.exists() else -1)
        
        fixes.append({'type': 'saved', 'message': f'已儲存至 {output_path}'})
        log.info("=== fix_document 結束 ===")
        return fixes
    
    def _fix_page_margins(self, doc) -> List[Dict]:
        margins = self.rule_engine.get_page_margins("body")
        for section in doc.sections:
            section.top_margin = Cm(margins[0])
            section.bottom_margin = Cm(margins[1])
            section.left_margin = Cm(margins[2])
            section.right_margin = Cm(margins[3])
        return [{'type': 'page_margin', 'message': f'已修正頁邊距'}]
    
    def _fix_heading_style_defs(self, doc) -> List[Dict]:
        """只修改樣式定義，不動個別段落"""
        heading_styles = self.rule_engine.get('heading_styles', {})
        style_map = {
            'heading1': 'Heading 1',
            'heading2': 'Heading 2',
            'heading3': 'Heading 3',
            'heading4': 'Heading 4',
        }
        align_map = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT,
                     'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
        
        count = 0
        for rule_key, style_name in style_map.items():
            config = heading_styles.get(rule_key, {})
            if not config:
                continue
            try:
                style = doc.styles[style_name]
                style.font.size = Pt(config.get('font_size', 12))
                style.font.bold = config.get('bold', True)
                # 設定中文字型
                style.font.name = config.get('font_name', '標楷體')
                # 設定英文字型為 Times New Roman
                rPr = style._element.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    style._element.append(rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), config.get('english_font', 'Times New Roman'))
                rFonts.set(qn('w:hAnsi'), config.get('english_font', 'Times New Roman'))
                rFonts.set(qn('w:eastAsia'), config.get('font_name', '標楷體'))
                pf = style.paragraph_format
                pf.alignment = align_map.get(config.get('align', 'left'))
                pf.space_before = Pt(config.get('before', 0))
                pf.space_after = Pt(config.get('after', 0))
                count += 1
            except Exception as e:
                log.warning("Heading %s 樣式設定失敗: %s", rule_key, e)
        
        # 設定 Title 樣式（摘要、Abstract、致謝、圖目錄、表目錄）為 20pt
        try:
            title_style = doc.styles['Title']
            title_cfg = heading_styles.get('title', {})
            title_font_size = title_cfg.get('font_size', 20)
            title_font_bold = title_cfg.get('bold', True)
            title_font_zh = title_cfg.get('font_name', self.rule_engine.get('fonts.chinese_title.name', '標楷體'))
            title_font_en = title_cfg.get('english_font', self.rule_engine.get('fonts.english.name', 'Times New Roman'))
            title_style.font.size = Pt(title_font_size)
            title_style.font.bold = title_font_bold
            rPr = title_style._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                title_style._element.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), title_font_en)
            rFonts.set(qn('w:hAnsi'), title_font_en)
            rFonts.set(qn('w:eastAsia'), title_font_zh)
            count += 1
        except Exception as e:
            log.warning("Title 樣式設定失敗: %s", e)
        
        body_cfg = self.rule_engine.get('body_style', {})
        body_font = body_cfg.get('font_name', '新細明體')
        body_eng = body_cfg.get('english_font', 'Times New Roman')
        body_align = body_cfg.get('align', 'justify')

        align_map_full = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT,
                          'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}

        # 設定 Normal 樣式（正文）
        try:
            normal_style = doc.styles['Normal']
            normal_style.font.size = Pt(body_cfg.get('font_size', 12))
            rPr = normal_style._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                normal_style._element.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), body_eng)
            rFonts.set(qn('w:hAnsi'), body_eng)
            rFonts.set(qn('w:eastAsia'), body_font)
            normal_style.paragraph_format.alignment = align_map_full.get(body_align)
            count += 1
        except Exception as e:
            log.warning("Normal 樣式設定失敗: %s", e)

        # 設定 Caption 樣式（圖說/表說）跟隨內文字體
        try:
            cap_style = doc.styles['Caption']
            cap_style.font.size = Pt(body_cfg.get('font_size', 12))
            rPr = cap_style._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                cap_style._element.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), body_eng)
            rFonts.set(qn('w:hAnsi'), body_eng)
            rFonts.set(qn('w:eastAsia'), body_font)
            cap_style.paragraph_format.alignment = align_map_full.get(body_align)
            count += 1
        except Exception as e:
            log.warning("Caption 樣式設定失敗: %s", e)
        
        return [{'type': 'heading_style', 'message': f'已更新{count}個標題樣式定義'}]

    def _remove_heading_indent(self, doc) -> List[Dict]:
        """移除 Heading 1 段落的首行縮排（直接格式）"""
        heading1_styles = {'Heading 1', '1'}
        count = 0
        for p in doc.paragraphs:
            pstyle = p.style.name if p.style else ''
            if pstyle not in heading1_styles:
                continue
            pPr = p._element.find(qn('w:pPr'))
            if pPr is None:
                continue
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                continue
            if ind.get(qn('w:firstLine')) is not None:
                del ind.attrib[qn('w:firstLine')]
                count += 1
        return [{'type': 'heading_indent', 'message': f'已移除{count}個Heading 1首行縮排'}]
    
    def _fix_heading_page_break(self, doc) -> List[Dict]:
        """為所有 Heading 1 設定 pageBreakBefore"""
        h1_paras = [p for p in doc.paragraphs if p.style.name == 'Heading 1' and p.text.strip()]
        
        count = 0
        for para in h1_paras:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            
            pb = pPr.find(qn('w:pageBreakBefore'))
            if pb is None:
                pb = OxmlElement('w:pageBreakBefore')
                pPr.append(pb)
            
            if pb.get(qn('w:val')) != '1':
                pb.set(qn('w:val'), '1')
                count += 1
        
        if count:
            return [{'type': 'page_break', 'message': f'已為{count}個章節標題設定分頁'}]
        return []
    
    def _fix_line_spacing(self, doc) -> List[Dict]:
        """修正內文行距為1.5倍（跳過表格）"""
        expected = self.rule_engine.get('line_spacing.value', 1.5)
        count = 0
        
        for para in doc.paragraphs:
            # 跳過表格內段落
            if para._element.getparent().tag.endswith('}tc'):
                continue
            
            pf = para.paragraph_format
            current = pf.line_spacing
            rule = pf.line_spacing_rule
            
            if (
                rule == WD_LINE_SPACING.ONE_POINT_FIVE
                or (
                    rule == WD_LINE_SPACING.MULTIPLE
                    and isinstance(current, (int, float))
                    and abs(current - expected) < 0.1
                )
            ):
                continue
            
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = expected
            count += 1
        
        if count:
            return [{'type': 'line_spacing', 'message': f'已修正{count}段行距為{expected}倍'}]
        return []
    
    def _fix_numbering_xml(self, doc) -> List[Dict]:
        """修改編號定義: "-"改"."，數字字型改TNR"""
        numbering_part = doc.part.numbering_part
        if not numbering_part:
            return []
        
        changed = 0
        for abs_num in numbering_part._element.findall(qn('w:abstractNum')):
            for lvl in abs_num.findall(qn('w:lvl')):
                txt_el = lvl.find(qn('w:lvlText'))
                if txt_el is None:
                    continue
                
                val = txt_el.get(qn('w:val'))
                if not val:
                    continue
                
                # 改 "-" 為 "."
                if '-' in val:
                    txt_el.set(qn('w:val'), val.replace('-', '.'))
                    changed += 1
                
                # 數字字型改 TNR
                rPr = lvl.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    lvl.append(rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        if changed:
            return [{'type': 'numbering', 'message': f'已修改{changed}個編號格式 ( "-" → "." )'}]
        return []
    
    def _fix_empty_paragraphs(self, doc) -> List[Dict]:
        """刪除多餘的空段落（只刪完全空的段落，保留有文字的段落與欄位碼段落）"""
        body = doc.element.body
        
        # 先找出所有 Heading 的位置
        heading_indices = set()
        all_children = list(body)
        for i, child in enumerate(all_children):
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    val = pStyle.get(qn('w:val'))
                    if val and 'Heading' in val:
                        heading_indices.add(i)
        
        # 收集要刪除的空段落（從後往前）
        to_remove = []
        for i, child in enumerate(all_children):
            if not child.tag.endswith('}p'):
                continue
            if i in heading_indices:
                continue
            # 跳過表格內段落
            parent_tag = child.getparent().tag if child.getparent() is not None else ''
            if 'tc' in parent_tag:
                continue
            # 跳過有欄位碼的段落（如 TOC、交叉參照）
            if child.findall('.//' + qn('w:fldChar')):
                continue
            # 跳過含 sectPr 的段落（分節設定，刪除會丟失 section）
            pPr_elem = child.find(qn('w:pPr'))
            if pPr_elem is not None and pPr_elem.find(qn('w:sectPr')) is not None:
                continue
            # 檢查是否為空段落 - 用文字內容判斷
            text_content = ''.join(t.text or '' for t in child.findall('.//' + qn('w:t')))
            if text_content.strip():
                continue  # 有文字，不刪除
            # 也不刪除有圖片或其他內容的段落
            if child.findall('.//' + qn('w:drawing')):
                continue
            if child.findall('.//' + qn('w:pict')):
                continue
            to_remove.append(child)
        
        for elem in to_remove:
            body.remove(elem)
        
        if to_remove:
            return [{'type': 'empty_paragraphs', 'message': f'已刪除{len(to_remove)}個空段落'}]
        return []
    
    def _fix_special_sections_page_break(self, doc) -> List[Dict]:
        """為特殊章節設定分頁（摘要、Abstract、目錄、圖目錄、表目錄、致謝）"""
        keywords = ['摘要', 'abstract', '目錄', '圖目錄', '表目錄', '致謝',
                     'acknowledgement', 'table of contents', 'list of figures', 'list of tables']
        
        count = 0
        
        # 處理一般段落
        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name
            
            # 檢查是否為特殊章節（僅限標題樣式）
            is_special = False
            if 'Heading' in style:
                for kw in keywords:
                    if text.lower() == kw or text == kw:
                        is_special = True
                        break
            elif style in ('Title', 'a9'):
                for kw in keywords:
                    lower_text = text.lower()
                    kw_lower = kw.lower()
                    if kw_lower in lower_text:
                        is_special = True
                        break
            
            if not is_special:
                continue
            
            # 設定 pageBreakBefore
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            
            pb = pPr.find(qn('w:pageBreakBefore'))
            if pb is None:
                pb = OxmlElement('w:pageBreakBefore')
                pPr.append(pb)
            
            if pb.get(qn('w:val')) != '1':
                pb.set(qn('w:val'), '1')
                count += 1
        
        # 處理 sdt 內的目錄（Word 生成的 TOC）
        body = doc.element.body
        for child in body:
            if not child.tag.endswith('}sdt'):
                continue
            sdtContent = child.find(qn('w:sdtContent'))
            if sdtContent is None:
                continue
            text = ''.join(t.text or '' for t in sdtContent.findall('.//' + qn('w:t')))
            if '目錄' not in text and '目錄' not in text:
                continue
            
            # 在 sdtContent 中找第一個段落，設定 pageBreakBefore
            first_para = sdtContent.find('.//' + qn('w:p'))
            if first_para is not None:
                pPr = first_para.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    first_para.insert(0, pPr)
                pb = pPr.find(qn('w:pageBreakBefore'))
                if pb is None:
                    pb = OxmlElement('w:pageBreakBefore')
                    pPr.append(pb)
                if pb.get(qn('w:val')) != '1':
                    pb.set(qn('w:val'), '1')
                    count += 1
        
        if count:
            return [{'type': 'special_section_page_break', 'message': f'已為{count}個特殊章節設定分頁'}]
        return []
    
    def _fix_toc_hyperlinks(self, doc) -> List[Dict]:
        """移除目錄超連結、rStyle=Hyperlink，並設定黑色文字"""
        count = 0
        rstyle_removed = 0
        body = doc.element.body
        
        def _fix_container(container):
            nonlocal count, rstyle_removed
            # 1) 移除所有 runs 上的 rStyle=Hyperlink (字元樣式)
            for run in container.iter(qn('w:r')):
                rPr = run.find(qn('w:rPr'))
                if rPr is not None:
                    rStyle = rPr.find(qn('w:rStyle'))
                    if rStyle is not None and rStyle.get(qn('w:val'), '').lower() == 'hyperlink':
                        rPr.remove(rStyle)
                        rstyle_removed += 1
                    color = rPr.find(qn('w:color'))
                    if color is not None:
                        color.set(qn('w:val'), '000000')
                    else:
                        color = OxmlElement('w:color')
                        color.set(qn('w:val'), '000000')
                        rPr.append(color)
                    u = rPr.find(qn('w:u'))
                    if u is not None:
                        rPr.remove(u)
            
            # 2) 移除 hyperlink 元素，內含 runs 保留原位
            for hl in list(container.iter(qn('w:hyperlink'))):
                parent = hl.getparent()
                if parent is None:
                    continue
                runs = hl.findall('.//' + qn('w:r'))
                idx = list(parent).index(hl)
                for run in reversed(runs):
                    parent.insert(idx, run)
                parent.remove(hl)
                count += 1
        
        # 處理 SDT 內容
        for sdt in body:
            if not sdt.tag.endswith('}sdt'):
                continue
            sdtContent = sdt.find(qn('w:sdtContent'))
            if sdtContent is None:
                continue
            _fix_container(sdtContent)
        
        # 處理一般 body 段落（含 圖目錄、表目錄條目）
        _fix_container(body)
        
        log.info("_fix_toc_hyperlinks: 移除 %d hyperlinks, 移除 %d rStyle=Hyperlink", count, rstyle_removed)
        
        # 3) 直接修改 Hyperlink 字元樣式定義
        try:
            hl_style = doc.styles['Hyperlink']
            hl_style.font.underline = False
            hl_style.font.color.rgb = None
            log.info("已清除 Hyperlink 樣式定義")
        except Exception as e:
            log.warning("找不到 Hyperlink 樣式: %s", e)
        
        if count or rstyle_removed:
            return [{'type': 'toc_hyperlinks', 'message': f'已移除{count}個超連結、清除{rstyle_removed}個Hyperlink樣式'}]
        return []
    
    @staticmethod
    def _make_xml_element(tag, **attrs):
        from lxml import etree
        NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        if tag.startswith('w:'):
            tag = tag[2:]
        el = etree.SubElement(etree.Element('dummy'), '{' + NS + '}' + tag)
        for k, v in attrs.items():
            if k.startswith('w:'):
                k = k[2:]
            el.set('{' + NS + '}' + k, str(v))
        return el
    
    # Word 自動產生的 TOC 縮排標準值（樣式定義層級）
    _TOC_STYLE_INDENTS = {
        'toc 1': {'firstLine': '0', 'firstLineChars': '0'},
        'toc 2': {'left': '200', 'leftChars': '200', 'firstLine': '0', 'firstLineChars': '0'},
        'toc 3': {'left': '400', 'leftChars': '400', 'firstLine': '0', 'firstLineChars': '0'},
    }
    # Word 自動產生的 TOC 段落直接格式縮排值
    _TOC_PARA_INDENTS = {
        'toc 2': {'left': '520'},
        'toc 3': {'left': '1040'},
    }
    # Word 自動產生的圖表目錄縮排（樣式定義層級）
    _TOF_STYLE_INDENTS = {'left': '100', 'hanging': '100', 'hangingChars': '100'}
    # Word 自動產生的圖表目錄段落直接格式縮排
    _TOF_PARA_INDENTS = {'left': '260', 'hanging': '260'}
    
    def _fix_toc_style(self, doc) -> List[Dict]:
        """修正 TOC & 圖表目錄樣式定義：字型、粗體、縮排、凸排"""
        count = 0
        bold_removed = 0
        rstyle_removed = 0
        indent_fixed = 0
        body_cfg = self.rule_engine.get('body_style', {})
        body_font = body_cfg.get('font_name', '新細明體')
        body_eng = body_cfg.get('english_font', 'Times New Roman')
        
        # 建立 name → (style_obj, style_id) 對照
        name_map = {}
        for s in doc.styles:
            nl = s.name.lower()
            sid = s._element.get(qn('w:styleId'))
            name_map[nl] = (s, sid)
        
        # ---- TOC 樣式：修正字型，設對齊靠左，設定定位點 ----
        for toc_name in ('toc 1', 'toc 2', 'toc 3'):
            entry = name_map.get(toc_name)
            if entry is None:
                continue
            style, sid = entry
            # 字型修正
            style.font.underline = False
            style.font.color.rgb = None
            rPr = style._element.find(qn('w:rPr'))
            if rPr is not None:
                u = rPr.find(qn('w:u'))
                if u is not None:
                    rPr.remove(u)
                color = rPr.find(qn('w:color'))
                if color is not None:
                    color.set(qn('w:val'), '000000')
                rStyle = rPr.find(qn('w:rStyle'))
                if rStyle is not None and rStyle.get(qn('w:val'), '').lower() == 'hyperlink':
                    rPr.remove(rStyle)
            # 設定字型
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style._element.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), body_eng)
            rFonts.set(qn('w:hAnsi'), body_eng)
            rFonts.set(qn('w:eastAsia'), body_font)
            # 設對齊靠左（避免 Normal justify 傳入）
            pPr = style._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                style._element.insert(0, pPr)
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'left')
            count += 1
        
        # ---- table of figures 樣式：清除粗體，修正字型 ----
        tof_entry = name_map.get('table of figures')
        if tof_entry is not None:
            style, sid = tof_entry
            style.font.bold = False
            count += 1
            # 設對齊靠左
            pPr = style._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                style._element.insert(0, pPr)
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'left')
            # 設定字型
            rPr = style._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style._element.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), body_eng)
            rFonts.set(qn('w:hAnsi'), body_eng)
            rFonts.set(qn('w:eastAsia'), body_font)
        
        # ---- 段落層級：只清除 run 粗體，保留原始縮排 ----
        for para in doc.element.body.iter(qn('w:p')):
            pPr = para.find(qn('w:pPr'))
            if pPr is None:
                continue
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                continue
            style_id = pStyle.get(qn('w:val'))
            if not style_id:
                continue
            is_toc = style_id in (sid for name_lower, (st, sid) in name_map.items() if name_lower in ('toc 1', 'toc 2', 'toc 3'))
            is_tof = style_id in (sid for name_lower, (st, sid) in name_map.items() if name_lower == 'table of figures')
            if not is_toc and not is_tof:
                continue
            # 清除 run 層級粗體 + 字元樣式 + 字型大小（讓樣式層級統一控制）
            for r in para.findall('.//' + qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is not None:
                    b = rPr.find(qn('w:b'))
                    if b is not None:
                        rPr.remove(b)
                        bold_removed += 1
                    rs_elem = rPr.find(qn('w:rStyle'))
                    if rs_elem is not None:
                        rPr.remove(rs_elem)
                        rstyle_removed += 1
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        rPr.remove(sz)
                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is not None:
                        rPr.remove(szCs)
        
        # ---- 更新 TOC 字型大小 ----
        paper_cfg = self.rule_engine.get('paper', {})
        toc_font_size = paper_cfg.get('toc_font_size', 12)
        for toc_name in ('toc 1', 'toc 2', 'toc 3', 'table of figures'):
            entry = name_map.get(toc_name)
            if entry is None:
                continue
            style, sid = entry
            rPr = style._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style._element.append(rPr)
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                rPr.remove(sz)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(toc_font_size * 2))
            rPr.append(sz)
            szCs = rPr.find(qn('w:szCs'))
            if szCs is not None:
                rPr.remove(szCs)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(toc_font_size * 2))
            rPr.append(szCs)

        # ---- 更新 TOC 右對齊定位點（隨頁寬調整） ----
        pw = paper_cfg.get('width_cm', 21.0)
        margins = self.rule_engine.get('page_margins', {}).get('body', {})
        ml = margins.get('left', 2.54)
        mr = margins.get('right', 2.54)
        avail_cm = pw - ml - mr
        tab_pos_twips = int(avail_cm / 2.54 * 1440)

        def _upsert_right_tab(pPr, pos_twips):
            """更新或加入右對齊定位點，保留其他定位點"""
            tabs = pPr.find(qn('w:tabs'))
            if tabs is None:
                tabs = OxmlElement('w:tabs')
                pPr.append(tabs)
            found = False
            for tab in tabs.findall(qn('w:tab')):
                val = tab.get(qn('w:val'))
                leader = tab.get(qn('w:leader'))
                if val == 'right' and leader == 'dot':
                    tab.set(qn('w:pos'), str(pos_twips))
                    found = True
                    break
            if not found:
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), str(pos_twips))
                tab.set(qn('w:leader'), 'dot')
                tabs.append(tab)

        # 更新 TOC / TOF 樣式定位點
        for toc_name in ('toc 1', 'toc 2', 'toc 3', 'table of figures'):
            entry = name_map.get(toc_name)
            if entry is None:
                continue
            style, sid = entry
            pPr = style._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                style._element.insert(0, pPr)
            _upsert_right_tab(pPr, tab_pos_twips)

        # 更新段落層級定位點
        target_ids = set(sid for n, (s, sid) in name_map.items() if n in ('toc 1', 'toc 2', 'toc 3', 'table of figures'))
        for para in doc.element.body.iter(qn('w:p')):
            pPr = para.find(qn('w:pPr'))
            if pPr is None:
                continue
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                continue
            style_id = pStyle.get(qn('w:val'))
            if not style_id:
                continue
            if style_id in target_ids:
                _upsert_right_tab(pPr, tab_pos_twips)

        log.info("_fix_toc_style: 修正 %d 個樣式/段落, 處理 %d 個縮排、%d 個粗體、%d 個字元樣式",
                 count, indent_fixed, bold_removed, rstyle_removed)
        
        if count or indent_fixed or bold_removed or rstyle_removed:
            parts = [f'已修正{count}個目錄樣式']
            if indent_fixed:
                parts.append(f'處理{indent_fixed}個縮排/凸排')
            if bold_removed:
                parts.append(f'移除{bold_removed}個粗體')
            if rstyle_removed:
                parts.append(f'移除{rstyle_removed}個字元樣式')
            parts.append(f'定位點→{avail_cm:.1f}cm')
            return [{'type': 'toc_style', 'message': '、'.join(parts)}]
        return []
    
    def _fix_caption_numbering(self, doc) -> List[Dict]:
        """修正圖表說明中的編號格式: 2-1 → 2.1（處理所有段落）"""
        CAPTION_RE = re.compile(r'((?:[图表圖表]|\b[Ff]ig(?:ure)?\.?))\s*(\d+)-(\d+)')
        count = 0
        body = doc.element.body
        para_matches = 0
        sdt_matches = 0
        
        def _replace_nbh_with_dot(parent_element):
            """將 noBreakHyphen 取代為含「.」的新 run"""
            c = 0
            for nbh in list(parent_element.findall('.//' + qn('w:noBreakHyphen'))):
                parent_run = nbh.getparent()
                if parent_run is not None and parent_run.tag != qn('w:r'):
                    continue
                if parent_run is not None:
                    grandparent = parent_run.getparent()
                    if grandparent is not None:
                        # 判斷是否為 Case A: prev run 結尾為數字且此 run 開頭為數字 → 插在中間
                        prev_r = parent_run.getprevious()
                        prev_text = ''
                        while prev_r is not None:
                            if prev_r.tag == qn('w:r'):
                                pt = ''.join(t.text or '' for t in prev_r.findall(qn('w:t')))
                                if pt:
                                    prev_text = pt
                                    break
                            prev_r = prev_r.getprevious()
                        this_text = ''.join(t.text or '' for t in parent_run.findall(qn('w:t')))
                        prev_stripped = prev_text.rstrip()
                        this_stripped = this_text.lstrip()
                        if (prev_stripped and prev_stripped[-1].isdigit() and
                                this_stripped and this_stripped[0].isdigit()):
                            new_r = OxmlElement('w:r')
                            new_t = OxmlElement('w:t')
                            new_t.text = '.'
                            new_r.append(new_t)
                            prev_r.addnext(new_r)
                        else:
                            new_r = OxmlElement('w:r')
                            new_t = OxmlElement('w:t')
                            new_t.text = '.'
                            new_r.append(new_t)
                            idx = list(grandparent).index(parent_run)
                            grandparent.insert(idx + 1, new_r)
                    parent_run.remove(nbh)
                    c += 1
            return c
        
        # 處理所有段落（不限 style）
        for para in doc.paragraphs:
            text = para.text
            if not CAPTION_RE.search(text):
                continue
            para_matches += 1
            
            # 取代 noBreakHyphen → 插入「.」
            count += _replace_nbh_with_dot(para._element)
            
            # 也處理普通連字號（-）
            for r in para._element.findall('.//' + qn('w:r')):
                t = r.find(qn('w:t'))
                if t is not None and t.text and '-' in t.text:
                    text_val = t.text
                    new_text = re.sub(r'(\d)-(\d)', r'\1.\2', text_val)
                    if new_text != text_val:
                        t.text = new_text
                        count += 1
        
        # 處理 sdtContent 內的目錄圖表條目
        for child in body:
            if not child.tag.endswith('}sdt'):
                continue
            sdtContent = child.find(qn('w:sdtContent'))
            if sdtContent is None:
                continue
            
            sdt_text = ''.join(t.text or '' for t in sdtContent.findall('.//' + qn('w:t')))
            if not CAPTION_RE.search(sdt_text):
                continue
            sdt_matches += 1
            
            count += _replace_nbh_with_dot(sdtContent)
            
            for r in sdtContent.findall('.//' + qn('w:r')):
                t = r.find(qn('w:t'))
                if t is not None and t.text and '-' in t.text:
                    new_text = re.sub(r'(\d)-(\d)', r'\1.\2', t.text)
                    if new_text != t.text:
                        t.text = new_text
                        count += 1
        
        # 第三遍：處理目錄條目中相鄰的章節/圖表數字 (如 "圖 21" → "圖 2.1")
        ADJACENT_RE = re.compile(r'((?:[图表圖表]|\b[Ff]ig(?:ure)?\.?))\s+(\d)(\d)(?:\s|\.|$)')
        for para in doc.paragraphs:
            m = ADJACENT_RE.search(para.text)
            if not m:
                continue
            ch_digit, fig_digit = m.group(2), m.group(3)
            # 跳過已有點的條目
            if '.' in m.group(0):
                continue
            runs = [child for child in para._element if child.tag == qn('w:r')]
            for i in range(len(runs) - 1):
                r1_text = ''.join(t.text or '' for t in runs[i].findall(qn('w:t')))
                r2_text = ''.join(t.text or '' for t in runs[i+1].findall(qn('w:t')))
                if r1_text.rstrip().endswith(ch_digit) and r2_text.lstrip().startswith(fig_digit):
                    new_r = OxmlElement('w:r')
                    new_t = OxmlElement('w:t')
                    new_t.text = '.'
                    new_r.append(new_t)
                    runs[i].addnext(new_r)
                    count += 1
                    break
        
        log.info("_fix_caption_numbering: %d 段落匹配, %d sdt 匹配, %d 處修正", para_matches, sdt_matches, count)
        
        if count:
            return [{'type': 'caption_numbering', 'message': f'已修正{count}處圖表編號格式 ( "-" → "." )'}]
        return []

    def _fix_heading_numbering(self, doc) -> List[Dict]:
        """將標題與目錄的編號格式 1-1 → 1.1"""
        cfg = self.rule_engine.get('heading_numbering_conversion', {})
        if not cfg.get('from_dash', True):
            return []
        heading_styles = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                          'toc 1', 'toc 2', 'toc 3', 'toc 4'}
        tgt_re = re.compile(r'(\d)-(\d)')
        count = 0

        def _fix_element(parent):
            c = 0
            for r in parent.findall('.//' + qn('w:r')):
                t = r.find(qn('w:t'))
                if t is not None and t.text and '-' in t.text:
                    new_text = tgt_re.sub(r'\1.\2', t.text)
                    if new_text != t.text:
                        t.text = new_text
                        c += 1
            return c

        # 處理一般段落
        for para in doc.paragraphs:
            sname = para.style.name if para.style else ''
            if sname not in heading_styles:
                continue
            if not para.text.strip():
                continue
            count += _fix_element(para._element)

        # 處理 SDT 包裝的目錄段落（樣式名可能是內部 ID 如 11, 21, 31）
        for child in doc.element.body:
            if not child.tag.endswith('}sdt'):
                continue
            sdt = child.find(qn('w:sdtContent'))
            if sdt is None:
                continue
            sdt_text = ''.join(t.text or '' for t in sdt.findall('.//' + qn('w:t')))
            if not tgt_re.search(sdt_text):
                continue
            count += _fix_element(sdt)

        if count:
            return [{'type': 'heading_numbering', 'message': f'已修正{count}處標題/目錄編號 ( "-" → "." )'}]
        return []

    def _fix_caption_alignment(self, doc) -> List[Dict]:
        """依配置設定圖表說明對齊方式"""
        ft_rules = self.rule_engine.get_figure_table_rules()
        fig_align = ft_rules.get('figures', {}).get('caption_align', 'center')
        tbl_align = ft_rules.get('tables', {}).get('caption_align', 'center')
        xml_align_map = {'center': 'center', 'justify': 'both', 'left': 'left', 'right': 'right'}
        CAPTION_RE = re.compile(r'^((?:[图表圖表]|\b[Ff]ig(?:ure)?\.?))\s+\d+\.\d+')
        count = 0
        for para in doc.paragraphs:
            m = CAPTION_RE.match(para.text.strip())
            if not m:
                continue
            prefix = m.group(1)
            align_key = tbl_align if prefix in '表' else fig_align
            xml_val = xml_align_map.get(align_key)
            if xml_val is None:
                continue
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), xml_val)
            count += 1
        if count:
            return [{'type': 'caption_alignment', 'message': f'已修正{count}處圖表說明對齊方式'}]
        return []

    def _fix_table_alignment(self, doc) -> List[Dict]:
        """設定所有表格置中對齊"""
        count = 0
        for table in doc.tables:
            tbl = table._element
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            jc = tblPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                tblPr.append(jc)
            jc.set(qn('w:val'), 'center')
            count += 1
        if count:
            return [{'type': 'table_alignment', 'message': f'已設定{count}個表格置中對齊'}]
        return []

    def _fix_table_no_break_across_pages(self, doc) -> List[Dict]:
        """禁止表格列跨頁切割"""
        count = 0
        for table in doc.tables:
            for row in table.rows:
                tr = row._element
                trPr = tr.find(qn('w:trPr'))
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    tr.insert(0, trPr)
                cantSplit = trPr.find(qn('w:cantSplit'))
                if cantSplit is None:
                    cantSplit = OxmlElement('w:cantSplit')
                    trPr.append(cantSplit)
                if cantSplit.get(qn('w:val')) in (None, 'true', '1', 'on'):
                    cantSplit.set(qn('w:val'), '1')
                    count += 1
        if count:
            return [{'type': 'table_no_break', 'message': f'已設定{count}個表格列禁止跨頁切割'}]
        return []

    def _fix_toc_field_codes(self, doc) -> List[Dict]:
        r"""移除 TOC \c 及 BIBLIOGRAPHY 欄位結構，結果文字保留為純文字；
        使用精確比對避免誤傷一般交叉參照"""
        body = doc.element.body
        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        removed = 0

        all_paras = list(body.iter('{' + W + '}p'))

        for p in all_paras:
            runs = list(p.findall(qn('w:r')))
            i = 0
            while i < len(runs):
                r = runs[i]
                fld = r.find(qn('w:fldChar'))
                if fld is not None and fld.get(qn('w:fldCharType')) == 'begin':
                    is_target = False
                    instr_text_accum = ''
                    for r2 in runs[i+1:]:
                        fld2 = r2.find(qn('w:fldChar'))
                        if fld2 is not None:
                            break
                        instr = r2.find(qn('w:instrText'))
                        if instr is not None and instr.text and instr.text.strip():
                            instr_text_accum += instr.text.strip() + ' '
                    # 精確比對：TOC \c "..." 或 BIBLIOGRAPHY
                    compiled = re.compile(r'(?:TOC\s+\\c\s)|(?:^BIBLIOGRAPHY\s)')
                    if compiled.search(instr_text_accum.strip()):
                        is_target = True

                    if is_target:
                        to_remove = [runs[i]]
                        for r2 in runs[i+1:]:
                            fld2 = r2.find(qn('w:fldChar'))
                            if fld2 is not None and fld2.get(qn('w:fldCharType')) == 'separate':
                                to_remove.append(r2)
                                break
                            to_remove.append(r2)
                        for r in to_remove:
                            p.remove(r)
                            removed += 1
                i += 1

        if removed:
            log.info("_fix_toc_field_codes: 移除 %d 個欄位標記", removed)
            return [{'type': 'toc_field_codes', 'message': f'已移除 {removed} 個 TOC\\c/BIBLIOGRAPHY 欄位標記，結果文字保留為純文字'}]
        return []
    
    def _process_crossrefs_precise(self, cit_matches, fig_matches, char_run, runs, text, _set_bold, FIGURE_RE):
        from copy import deepcopy
        from docx.oxml.ns import qn

        run_bold = {}
        for m in cit_matches:
            for pos in range(m.start() + 1, m.end() - 1):
                ri, ci = char_run[pos]
                run_bold.setdefault(ri, set()).add(ci)
        for m in fig_matches:
            for pos in range(m.start(), m.end()):
                ri, ci = char_run[pos]
                run_bold.setdefault(ri, set()).add(ci)

        run_split = {}
        for m in cit_matches + fig_matches:
            for boundary in (m.start(), m.end()):
                if 0 < boundary < len(text):
                    ri, ci = char_run[boundary]
                    run_split.setdefault(ri, set()).add(ci)

        for ri in sorted(set(list(run_bold.keys()) + list(run_split.keys())), reverse=True):
            r = runs[ri]
            t_el = r.find(qn('w:t'))
            if t_el is None or t_el.text is None:
                continue
            rtxt = t_el.text
            bold_set = run_bold.get(ri, set())
            split_set = run_split.get(ri, set())

            if len(bold_set) == len(rtxt) and not split_set - {0, len(rtxt)}:
                _set_bold(r)
                continue
            if not bold_set and not split_set:
                continue

            parts = []
            prev = 0
            for off in sorted(split_set):
                if 0 < off < len(rtxt) and off > prev:
                    parts.append(rtxt[prev:off])
                    prev = off
            if prev < len(rtxt):
                parts.append(rtxt[prev:])

            if len(parts) <= 1:
                if bold_set:
                    _set_bold(r)
                continue

            parent = r.getparent()
            idx = list(parent).index(r)

            pos = 0
            for pt in reversed(parts):
                nr = deepcopy(r)
                nt = nr.find(qn('w:t'))
                if nt is not None:
                    nt.text = pt
                piece_bold = any(
                    bold_set.intersection(range(pos, pos + len(pt)))
                )
                if piece_bold:
                    _set_bold(nr)
                parent.insert(idx, nr)
                pos += len(pt)

            parent.remove(r)

    def _bold_figref_fallback(self, match, runs, parent_elem, _set_bold, FIGURE_RE):
        import re
        from docx.oxml.ns import qn

        fig_text = match.group()
        # First try: find a single run containing the full figure text
        for r in runs:
            t_el = r.find(qn('w:t'))
            if t_el is None or t_el.text is None:
                continue
            rtxt = t_el.text
            if fig_text in rtxt:
                _set_bold(r)
                return True
        
        # Second try: find runs that contain parts of the figure text
        # and bold them if they're adjacent and form the figure reference
        fig_parts = []
        for r in runs:
            t_el = r.find(qn('w:t'))
            if t_el is not None and t_el.text:
                rtxt = t_el.text
                # Check if this run contains any part of the figure text
                # Look for the specific figure pattern parts
                if re.search(r'(?:[圖表]|\b[Ff]ig(?:ure)?\.?)', rtxt) or re.search(r'\d+[-.]\d+', rtxt):
                    fig_parts.append(r)
        
        # If we found potential figure parts, bold them
        if fig_parts and len(fig_parts) >= 1:
            for r in fig_parts:
                _set_bold(r)
            return True
        
        return False

    def _fix_cross_references(self, doc) -> List[Dict]:
        """將交叉引用 [n] 中的數字設為粗體，圖表編號 圖2.1/表2.1 設為粗體"""
        import re
        from copy import deepcopy
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        CITATION_RE = re.compile(r'\[(\d+(?:[-,.]\d+)*)\]')
        FIGURE_RE = re.compile(r'((?:[圖表]|\b[Ff]ig(?:ure)?\.?))\s*(\d+[-.]\d+)')

        n_citation = 0
        n_figure = 0

        def _set_bold(run):
            rPr = run.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run.insert(0, rPr)
            b = rPr.find(qn('w:b'))
            if b is None:
                b = OxmlElement('w:b')
                rPr.append(b)
                return True
            return False

        for para in doc.paragraphs:
            text = para.text
            if not text:
                continue

            cit_matches = list(CITATION_RE.finditer(text))
            fig_matches = list(FIGURE_RE.finditer(text))
            if not cit_matches and not fig_matches:
                continue

            runs = list(para._element.findall(qn('w:r')))
            if not runs:
                continue

            # Build char-to-run mapping
            char_run = []
            for ri, r in enumerate(runs):
                t_el = r.find(qn('w:t'))
                rtxt = t_el.text if t_el is not None else ''
                for ci in range(len(rtxt)):
                    char_run.append((ri, ci))

            if len(char_run) == len(text):
                # Exact char-to-run mapping — handle both citations and figures
                self._process_crossrefs_precise(
                    cit_matches, fig_matches, char_run, runs, text,
                    _set_bold, FIGURE_RE
                )
                n_citation += len(cit_matches)
                n_figure += len(fig_matches)
            elif fig_matches:
                # Fallback: char mismatch (field codes) — check descendant runs
                desc_runs = list(para._element.findall('.//' + qn('w:r')))
                for m in fig_matches:
                    if self._bold_figref_fallback(m, desc_runs, para._element, _set_bold, FIGURE_RE):
                        n_figure += 1
                        # Debug: verify bold was applied
                        fig_text = m.group()
                        for r in desc_runs:
                            t_el = r.find(qn('w:t'))
                            if t_el is not None and t_el.text:
                                rtxt = t_el.text
                                if re.search(r'(?:[圖表]|\b[Ff]ig(?:ure)?\.?)', rtxt) or re.search(r'\d+[-.]\d+', rtxt):
                                    rPr = r.find(qn('w:rPr'))
                                    if rPr is not None:
                                        b = rPr.find(qn('w:b'))
                                        if b is not None:
                                            pass  # Bold confirmed

        total = n_citation + n_figure
        if total:
            return [{'type': 'cross_reference',
                     'message': f'已將 {n_citation} 處引用編號 [n] 及 {n_figure} 處圖表編號設為粗體'}]
        return []

    def _create_ref_numbering(self, doc, indent_val: int = 432) -> int:
        """建立 Word 自動編號定義 [%1]，回傳 numId"""
        from docx.oxml import OxmlElement

        numbering_part = doc.part.numbering_part
        numbering = numbering_part._element

        max_abstract = -1
        for an in numbering.findall(qn('w:abstractNum')):
            aid = int(an.get(qn('w:abstractNumId'), '0'))
            max_abstract = max(max_abstract, aid)
        new_abstract_id = max_abstract + 1

        max_num = 0
        for n_elem in numbering.findall(qn('w:num')):
            nid = int(n_elem.get(qn('w:numId'), '0'))
            max_num = max(max_num, nid)
        new_num_id = max_num + 1

        an = OxmlElement('w:abstractNum')
        an.set(qn('w:abstractNumId'), str(new_abstract_id))

        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'decimal')
        lvl.append(numFmt)

        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), '[%1]')
        lvl.append(lvlText)

        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '24')
        rPr.append(szCs)
        lvl.append(rPr)

        lvl_pPr = OxmlElement('w:pPr')
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:pos'), str(indent_val))
        tabs.append(tab)
        lvl_pPr.append(tabs)
        lvl_ind = OxmlElement('w:ind')
        lvl_ind.set(qn('w:left'), str(indent_val))
        lvl_ind.set(qn('w:hanging'), str(indent_val))
        lvl_ind.set(qn('w:firstLine'), '0')
        lvl_ind.set(qn('w:firstLineChars'), '0')
        lvl_pPr.append(lvl_ind)
        lvl.append(lvl_pPr)

        an.append(lvl)
        numbering.insert(0, an)

        num_elem = OxmlElement('w:num')
        num_elem.set(qn('w:numId'), str(new_num_id))
        an_ref = OxmlElement('w:abstractNumId')
        an_ref.set(qn('w:val'), str(new_abstract_id))
        num_elem.append(an_ref)
        numbering.append(num_elem)

        return new_num_id

    def _convert_references_to_numbered_list(self, doc) -> List[Dict]:
        """將參考文獻表格轉為 Word auto-numbering 段落，並加入 bookmark"""
        from docx.oxml import OxmlElement
        from copy import deepcopy

        def _is_ref_table(table):
            if len(table.columns) != 2:
                return False
            for row in table.rows:
                cell0_text = row.cells[0].text.strip()
                if not re.match(r'^\[\d+\]$', cell0_text):
                    return False
            return True

        ref_table = None
        for table in doc.tables:
            if _is_ref_table(table):
                ref_table = table
                break

        # 掃描現有 bookmark 以取得唯一 ID
        def _get_max_bookmark_id():
            max_bid = 0
            for p in doc.paragraphs:
                for bm in p._element.findall('.//' + qn('w:bookmarkStart')):
                    bid = int(bm.get(qn('w:id'), '0'))
                    max_bid = max(max_bid, bid)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for bm in p._element.findall('.//' + qn('w:bookmarkStart')):
                                bid = int(bm.get(qn('w:id'), '0'))
                                max_bid = max(max_bid, bid)
            return max_bid

        if ref_table is None:
            # 段落式參考文獻（1. 2. 3. 或 [1] [2] [3]）
            ref_heading_idx = None
            for i, p in enumerate(doc.paragraphs):
                txt = p.text.strip()
                if txt in ('參考文獻', '参考文献', 'REFERENCES', 'References') and p.style and 'heading' in p.style.name.lower():
                    ref_heading_idx = i
                    break
            if ref_heading_idx is None:
                return []

            bookmark_id = _get_max_bookmark_id() + 1
            count = 0
            REF_PARA_RE = re.compile(r'^(?:(\d+)\.\s*|\[(\d+)\]\s*)(.*)')
            for p in doc.paragraphs[ref_heading_idx + 1:]:
                txt = p.text.strip()
                if not txt:
                    continue
                m = REF_PARA_RE.match(txt)
                if not m:
                    continue
                ref_num = int(m.group(1) or m.group(2))
                # 刪除舊有同名 bookmark（避免重複）
                existing = p._element.findall('.//' + qn('w:bookmarkStart'))
                for bm in existing:
                    if bm.get(qn('w:name'), '').startswith('ref_'):
                        p._element.remove(bm)
                existing_end = p._element.findall('.//' + qn('w:bookmarkEnd'))
                for bm in existing_end:
                    if bm.get(qn('w:name'), '').startswith('ref_'):
                        p._element.remove(bm)
                bk_start = OxmlElement('w:bookmarkStart')
                bk_start.set(qn('w:id'), str(bookmark_id))
                bk_start.set(qn('w:name'), f'ref_{ref_num:03d}')
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    pPr.addnext(bk_start)
                else:
                    p._element.insert(0, bk_start)
                bk_end = OxmlElement('w:bookmarkEnd')
                bk_end.set(qn('w:id'), str(bookmark_id))
                p._element.append(bk_end)
                bookmark_id += 1
                count += 1
            if count:
                return [{'type': 'references_numbered_list',
                         'message': f'已為段落式參考文獻 {count} 條加入 bookmark'}]
            return []

        bookmark_id = _get_max_bookmark_id() + 1

        extracted = []
        ref_count = 0
        for row in ref_table.rows:
            num_text = row.cells[0].text.strip()
            num_match = re.match(r'^\[(\d+)\]$', num_text)
            if not num_match:
                continue
            ref_num = int(num_match.group(1))
            ref_para = row.cells[1].paragraphs[0]
            extracted.append((ref_num, ref_para))
            ref_count += 1

        extracted.sort(key=lambda x: x[0])

        indent_val = 600 if ref_count >= 10 else 432
        num_id = self._create_ref_numbering(doc, indent_val)

        new_p_elements = []
        for ref_num, orig_para in extracted:
            p_elem = OxmlElement('w:p')

            pPr = OxmlElement('w:pPr')

            numPr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numPr.append(ilvl)
            nid_el = OxmlElement('w:numId')
            nid_el.set(qn('w:val'), str(num_id))
            numPr.append(nid_el)
            pPr.append(numPr)

            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'left')
            pPr.append(jc)

            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(indent_val))
            ind.set(qn('w:hanging'), str(indent_val))
            ind.set(qn('w:firstLine'), '0')
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:leftChars'), '0')
            ind.set(qn('w:hangingChars'), '0')
            pPr.append(ind)

            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '360')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)

            p_elem.append(pPr)

            bk_start = OxmlElement('w:bookmarkStart')
            bk_start.set(qn('w:id'), str(bookmark_id))
            bk_start.set(qn('w:name'), f'ref_{ref_num:03d}')
            p_elem.append(bk_start)

            for child in orig_para._element:
                if child.tag == qn('w:r'):
                    p_elem.append(deepcopy(child))

            bk_end = OxmlElement('w:bookmarkEnd')
            bk_end.set(qn('w:id'), str(bookmark_id))
            p_elem.append(bk_end)

            new_p_elements.append(p_elem)
            bookmark_id += 1

        # 插入段落到表格原本的位置（addnext 插在後方，反序處理）
        parent = ref_table._element.getparent()
        table_elem = ref_table._element
        prev = table_elem
        for p_elem in new_p_elements:
            prev.addnext(p_elem)
            prev = p_elem

        parent.remove(table_elem)

        return [{'type': 'references_numbered_list',
                 'message': f'已將參考文獻表格轉為 {len(new_p_elements)} 條 auto-numbering 段落 + bookmark'}]

    def _convert_inline_citations_to_xrefs(self, doc) -> List[Dict]:
        """將內文手打的 [n] 引用轉為 REF field 交叉參照（指向參考文獻 bookmark）

        Surgical replacement：只替換含有 [n] 的純文字 runs，
        保留 field code runs（STYLEREF, SEQ, REF 等）完整不動。
        """
        from docx.oxml import OxmlElement
        from copy import deepcopy

        CITATION_RE = re.compile(r'\[(\d+(?:[-,.]\d+)*)\]')

        # 檢查文件中有無 ref_nnn bookmark，若無則跳過轉換
        has_ref_bookmarks = False
        for p in doc.paragraphs:
            for bm in p._element.findall('.//' + qn('w:bookmarkStart')):
                name = bm.get(qn('w:name'), '')
                if name.startswith('ref_'):
                    has_ref_bookmarks = True
                    break
            if has_ref_bookmarks:
                break
        if not has_ref_bookmarks:
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for bm in p._element.findall('.//' + qn('w:bookmarkStart')):
                                name = bm.get(qn('w:name'), '')
                                if name.startswith('ref_'):
                                    has_ref_bookmarks = True
                                    break
                        if has_ref_bookmarks:
                            break
                    if has_ref_bookmarks:
                        break
                if has_ref_bookmarks:
                    break
        if not has_ref_bookmarks:
            return []

        n_converted = 0

        ref_heading_pos = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == '參考文獻' and p.style and 'heading' in p.style.name.lower():
                ref_heading_pos = i
                break

        def _has_field_code(r_elem):
            return (r_elem.find(qn('w:fldChar')) is not None
                    or r_elem.find(qn('w:instrText')) is not None)

        def _make_text_run(text_val, bold=False):
            run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if bold:
                b_el = OxmlElement('w:b')
                rPr.append(b_el)
            run.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text_val
            run.append(t)
            return run

        def _make_ref_field(part_num):
            elements = []
            r1 = OxmlElement('w:r')
            rPr1 = OxmlElement('w:rPr')
            rPr1.append(OxmlElement('w:b'))
            r1.append(rPr1)
            fc1 = OxmlElement('w:fldChar')
            fc1.set(qn('w:fldCharType'), 'begin')
            r1.append(fc1)
            elements.append(r1)

            r2 = OxmlElement('w:r')
            instrT = OxmlElement('w:instrText')
            instrT.set(qn('xml:space'), 'preserve')
            instrT.text = f' REF ref_{int(part_num):03d} \\n \\h '
            r2.append(instrT)
            elements.append(r2)

            r3 = OxmlElement('w:r')
            fc_sep = OxmlElement('w:fldChar')
            fc_sep.set(qn('w:fldCharType'), 'separate')
            r3.append(fc_sep)
            elements.append(r3)

            r4 = OxmlElement('w:r')
            rPr4 = OxmlElement('w:rPr')
            rPr4.append(OxmlElement('w:b'))
            r4.append(rPr4)
            t4 = OxmlElement('w:t')
            t4.text = part_num
            r4.append(t4)
            elements.append(r4)

            r5 = OxmlElement('w:r')
            fc_end = OxmlElement('w:fldChar')
            fc_end.set(qn('w:fldCharType'), 'end')
            r5.append(fc_end)
            elements.append(r5)

            return elements

        for para in doc.paragraphs:
            if ref_heading_pos is not None:
                try:
                    p_idx = list(doc.element.body).index(para._element)
                    ref_elem_idx = list(doc.element.body).index(
                        doc.paragraphs[ref_heading_pos]._element)
                    if p_idx >= ref_elem_idx:
                        continue
                except ValueError:
                    pass

            all_runs = list(para._element.findall(qn('w:r')))
            if not all_runs:
                continue

            text_run_info = []
            for ri, r in enumerate(all_runs):
                if _has_field_code(r):
                    continue
                t_el = r.find(qn('w:t'))
                if t_el is not None and t_el.text:
                    text_run_info.append((ri, r, t_el.text))

            if not text_run_info:
                continue

            merged_text = ''.join(t for _, _, t in text_run_info)
            cit_matches = list(CITATION_RE.finditer(merged_text))
            if not cit_matches:
                continue

            tr_offsets = []
            off = 0
            for _, _, t in text_run_info:
                tr_offsets.append(off)
                off += len(t)

            # Collect all affected text-run indices and build replacement map
            # For each citation, identify which text-runs overlap and compute
            # the replacement elements (REF fields + brackets + separators)
            # along with the pre/post text fragments of partially-overlapping runs.

            # We walk through text_runs sequentially, consuming citation matches.
            # Strategy: replace affected text-runs with spliced elements.

            # Build a list of "segments" on the merged_text timeline:
            # each segment is either plain text (outside citations) or a citation.
            # Then reconstruct: remove all affected text-runs, insert new elements
            # at the position of the first removed run.

            segments = []
            last_seg_end = 0
            for m in cit_matches:
                if m.start() > last_seg_end:
                    segments.append(('text', merged_text[last_seg_end:m.start()],
                                     last_seg_end, m.start()))
                nums_str = m.group(1)
                segments.append(('citation', nums_str, m.start(), m.end()))
                last_seg_end = m.end()
                n_converted += 1
            if last_seg_end < len(merged_text):
                segments.append(('text', merged_text[last_seg_end:], last_seg_end,
                                 len(merged_text)))

            # Map merged_text character ranges to text_run indices
            # For each text segment, find which text_runs it spans
            # For each citation segment, build REF field elements

            def build_citation_elements(nums_str):
                elems = [_make_text_run('[', bold=True)]
                parts = [p.strip() for p in nums_str.split(',')]
                for pi, part in enumerate(parts):
                    sep = '-' if '-' in part else ('.' if '.' in part else None)
                    if sep:
                        start_num, end_num = part.split(sep, 1)
                        elems.extend(_make_ref_field(start_num))
                        elems.append(_make_text_run('-', bold=True))
                        elems.extend(_make_ref_field(end_num))
                    else:
                        elems.extend(_make_ref_field(part))
                    if pi < len(parts) - 1:
                        elems.append(_make_text_run(', ', bold=True))
                elems.append(_make_text_run(']', bold=True))
                return elems

            # Find which text_run indices are affected (fully or partially inside a citation)
            affected_tr_indices = set()
            for seg_type, seg_data, seg_start, seg_end in segments:
                if seg_type == 'citation':
                    for ti, (ri, r, t) in enumerate(text_run_info):
                        tr_start = tr_offsets[ti]
                        tr_end = tr_start + len(t)
                        if tr_start < seg_end and tr_end > seg_start:
                            affected_tr_indices.add(ti)

            if not affected_tr_indices:
                continue

            # Now reconstruct: walk through ALL children of the paragraph.
            # For non-text-run children (pPr, bookmarkStart/End, field runs, etc): keep as-is.
            # For text-run children:
            #   - If not affected: keep as-is
            #   - If affected: replace with rebuilt content

            # Simpler approach: remove affected text-runs, then insert new elements
            # at the position where the first affected run was.

            first_affected_tr = min(affected_tr_indices)
            first_affected_run = all_runs[text_run_info[first_affected_tr][0]]

            # Build all new elements for the affected region
            # We need to handle partial runs: a text-run may have text before/after
            # the citation region. We split those.

            # Collect all text_runs in order, marking which are affected
            # For affected runs, their full text is part of the "affected region"
            # but may include non-citation text that needs to be preserved.

            # Actually, the cleanest approach:
            # The merged_text for ALL text_runs is the concatenation.
            # The segments tell us exactly what goes where on this merged text.
            # We remove ALL affected text-runs and replace with new elements
            # that encode the segment content for the affected character range.

            # Find the character range in merged_text that covers all affected runs
            min_char = min(tr_offsets[ti] for ti in affected_tr_indices)
            max_char_ti = max(affected_tr_indices)
            max_char = tr_offsets[max_char_ti] + len(text_run_info[max_char_ti][2])

            # Build new elements for the affected character range
            new_elems = []
            for seg_type, seg_data, seg_start, seg_end in segments:
                # Clip segment to affected range
                clip_start = max(seg_start, min_char)
                clip_end = min(seg_end, max_char)
                if clip_start >= clip_end:
                    continue

                if seg_type == 'text':
                    # This text is within the affected range but outside any citation
                    # Extract the relevant substring
                    text_in_range = seg_data[clip_start - seg_start:clip_end - seg_start]
                    if text_in_range:
                        new_elems.append(_make_text_run(text_in_range))
                elif seg_type == 'citation':
                    new_elems.extend(build_citation_elements(seg_data))

            # Remove affected text-runs (in reverse to preserve indices)
            for ti in sorted(affected_tr_indices, reverse=True):
                ri = text_run_info[ti][0]
                para._element.remove(all_runs[ri])

            # Insert new elements at the position of the first removed run
            # Find the run that was just before the first affected run in the original
            first_ri = text_run_info[first_affected_tr][0]
            if first_ri > 0:
                # Find the nearest preceding sibling that still exists
                prev_run = None
                for ri_check in range(first_ri - 1, -1, -1):
                    if all_runs[ri_check] in list(para._element.findall(qn('w:r'))):
                        prev_run = all_runs[ri_check]
                        break

                if prev_run is not None:
                    for elem in reversed(new_elems):
                        prev_run.addnext(elem)
                else:
                    pPr = para._element.find(qn('w:pPr'))
                    if pPr is not None:
                        for elem in reversed(new_elems):
                            pPr.addnext(elem)
                    else:
                        for elem in reversed(new_elems):
                            para._element.insert(0, elem)
            else:
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    for elem in reversed(new_elems):
                        pPr.addnext(elem)
                else:
                    for elem in reversed(new_elems):
                        para._element.insert(0, elem)

        if n_converted:
            return [{'type': 'inline_citations_xref',
                      'message': f'已將 {n_converted} 處內文引用 [n] 轉為交叉參照 REF field'}]
        return []

    def _fix_direct_font_formatting(self, doc) -> List[Dict]:
        """移除段落 run 層級直接字型設定（rFonts/sz/szCs），保留粗體/斜體/底線；
        跳過欄位碼（fldChar/instrText）、頁首頁尾、目錄以避免誤傷"""
        count = 0
        # 收集所有頁首頁尾元素供跳過判斷
        hf_elements = set()
        for sec in doc.sections:
            for hf in (sec.header, sec.footer):
                hf_elements.add(hf._element)
        for para in doc.paragraphs:
            # 跳過頁首頁尾內的段落
            if para._element.getparent() in hf_elements:
                continue
            # 跳過目錄頁段落（含標題及內容行）
            txt = para.text.strip()
            sname = para.style.name if para.style else ''
            if sname in ('table of figures', 'table of tables', 'toc') or '目錄' in txt or 'Table of Contents' in txt:
                continue
            for run in para._element.findall(qn('w:r')):
                if run.find(qn('w:fldChar')) is not None or run.find(qn('w:instrText')) is not None:
                    continue  # 跳過欄位碼
                rPr = run.find(qn('w:rPr'))
                if rPr is None:
                    continue
                modified = False
                for tag in ['w:rFonts', 'w:sz', 'w:szCs']:
                    elem = rPr.find(qn(tag))
                    if elem is not None:
                        rPr.remove(elem)
                        modified = True
                if modified:
                    count += 1
        if count:
            return [{'type': 'direct_font', 'message': f'已清除{count}個段落直接字型設定'}]
        return []

    def _fix_first_line_indent(self, doc) -> List[Dict]:
        """為內文段落設定首行縮排（跳過標題、表格、封面、參考文獻）；
        數字清單改為凸排（左縮排=凸排量，首行=0）"""
        first_line_cfg = self.rule_engine.get('first_line_indent', {})
        if not first_line_cfg.get('enabled', True):
            return []
        indent_cm = first_line_cfg.get('value_cm', 0.75)
        skip_styles = {
            'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title',
            'table of figures', 'table of tables', 'toc',
        }
        cover_kw = self.rule_engine.get('cover.detect_keywords', ['雷射', 'Study on', 'MAR'])

        nl_cfg = self.rule_engine.get('numbered_list', {})
        nl_pattern = nl_cfg.get('pattern', r'^\d+[.、）)]')
        nl_hanging = nl_cfg.get('hanging_indent', True)
        nl_left = nl_cfg.get('hanging_indent_twips', 360)

        count_normal = 0
        count_list = 0
        in_cover = True

        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt in ('摘要', 'Abstract', '致謝', '目錄', '圖目錄', '表目錄', '參考文獻', '目次'):
                in_cover = False
            if in_cover:
                continue
            if para._element.getparent().tag.endswith('}tc'):
                continue
            style_name = para.style.name if para.style else ''
            if style_name in skip_styles:
                continue
            if not txt:
                continue

            # 圖表說跳過首行縮排（用 - 或 . 都配，因 caption_numbering 尚未執行）
            if re.match(r'^(?:[图表圖表]|\b[Ff]ig(?:ure)?\.?)\s*\d+[-.]\d+', txt):
                continue
            # 關鍵詞段落跳過首行縮排
            if re.match(r'^(關鍵詞|關鍵字|Key words|Keywords)\s*[:：]?\s*', txt, re.IGNORECASE):
                continue

            # 數字清單（手打 or Word 自動編號 w:numPr）→ 凸排
            pPr = para._element.find(qn('w:pPr'))
            has_numPr = pPr is not None and pPr.find(qn('w:numPr')) is not None
            if nl_hanging and (has_numPr or re.match(nl_pattern, txt)):
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    para._element.insert(0, pPr)
                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                ind.set(qn('w:left'), str(nl_left))
                ind.set(qn('w:hanging'), str(nl_left))
                ind.set(qn('w:firstLine'), '0')
                count_list += 1
                continue

            # 一般內文 → 首行縮排
            pf = para.paragraph_format
            current = pf.first_line_indent
            expected = Cm(indent_cm)
            if current is not None and abs(current.emu - expected.emu) < expected.emu * 0.1:
                continue
            pf.first_line_indent = Cm(indent_cm)
            count_normal += 1

        msgs = []
        if count_normal:
            msgs.append(f'已為{count_normal}段落設首行縮排 {indent_cm}cm')
        if count_list:
            msgs.append(f'已為{count_list}數字清單段落設凸排')
        if msgs:
            return [{'type': 'first_line_indent', 'message': '；'.join(msgs)}]
        return []

    def _fix_keywords_format(self, doc) -> List[Dict]:
        """將關鍵詞段落設為：無縮排、左右對齊、前空一行"""
        KW_RE = re.compile(r'^(關鍵詞|關鍵字|Key words|Keywords)\s*[:：]?\s*', re.IGNORECASE)
        count = 0
        for para in doc.paragraphs:
            txt = para.text.strip()
            if not KW_RE.match(txt):
                continue
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:firstLine'), '0')
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'both')
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            cur_before = spacing.get(qn('w:before'))
            if not cur_before or int(cur_before) < 240:
                spacing.set(qn('w:before'), '240')
            count += 1
        if count:
            return [{'type': 'keywords_format', 'message': f'已修正{count}處關鍵詞格式（無縮排、左右對齊、前空一行）'}]
        return []

    def _fix_page_numbers(self, doc) -> List[Dict]:
        """為各節設定頁碼（前言:羅馬數字, 正文:阿拉伯數字）"""
        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        pn_config = self.rule_engine.get('page_number', {})
        count = 0
        body_start_set = False
        for i, section in enumerate(doc.sections):
            is_front = (i == 0)
            config = pn_config.get('front_matter' if is_front else 'body', {})
            footer = section.footer
            footer.is_linked_to_previous = False
            ftr_elem = footer._element
            # 清除 footer 所有子元素（含 SDT、舊段落），只留一個乾淨的段落
            children = list(ftr_elem)
            for child in children:
                ftr_elem.remove(child)
            para_elem = OxmlElement('w:p')
            ftr_elem.append(para_elem)
            # 設定段落對齊
            align_map = {'center': 'center', 'left': 'left', 'right': 'right'}
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), align_map.get(config.get('align', 'center'), 'center'))
            pPr.append(jc)
            para_elem.insert(0, pPr)
            # 新增 PAGE 欄位
            fld_begin = OxmlElement('w:r')
            fc1 = OxmlElement('w:fldChar')
            fc1.set(qn('w:fldCharType'), 'begin')
            fld_begin.append(fc1)
            para_elem.append(fld_begin)
            fld_instr = OxmlElement('w:r')
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            fld_style = config.get('style', 'arabic')
            style_upper = 'ROMAN' if fld_style == 'roman_upper' else 'roman' if fld_style == 'roman_lower' else 'ARABIC'
            instr.text = f' PAGE \\* {style_upper}'
            fld_instr.append(instr)
            para_elem.append(fld_instr)
            fld_sep = OxmlElement('w:r')
            fc_sep = OxmlElement('w:fldChar')
            fc_sep.set(qn('w:fldCharType'), 'separate')
            fld_sep.append(fc_sep)
            para_elem.append(fld_sep)
            fld_end = OxmlElement('w:r')
            fc_end = OxmlElement('w:fldChar')
            fc_end.set(qn('w:fldCharType'), 'end')
            fld_end.append(fc_end)
            para_elem.append(fld_end)
            # 字型大小設定
            font_size = config.get('font_size', 12)
            for r in para_elem.findall(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)
                sz = rPr.find(qn('w:sz'))
                if sz is None:
                    sz = OxmlElement('w:sz')
                    rPr.append(sz)
                sz.set(qn('w:val'), str(font_size * 2))
                szCs = rPr.find(qn('w:szCs'))
                if szCs is None:
                    szCs = OxmlElement('w:szCs')
                    rPr.append(szCs)
                szCs.set(qn('w:val'), str(font_size * 2))
            # 設定頁碼起始值
            start_from = config.get('start_from')
            if start_from is not None:
                if is_front or not body_start_set:
                    sectPr = section._sectPr
                    pgNumType = sectPr.find(qn('w:pgNumType'))
                    if pgNumType is None:
                        pgNumType = OxmlElement('w:pgNumType')
                        sectPr.append(pgNumType)
                    pgNumType.set(qn('w:start'), str(start_from))
                    if not is_front:
                        body_start_set = True
            count += 1
        if count:
            return [{'type': 'page_number', 'message': f'已設定{count}節頁碼'}]
        return []

    def _fix_cover_tabs(self, doc) -> List[Dict]:
        """封面頁：調整各區塊間距"""
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        count = 0
        body_cfg = self.rule_engine.get('body_style', {})
        body_font = body_cfg.get('font_name', '新細明體')
        body_eng = body_cfg.get('english_font', 'Times New Roman')
        
        # 找封面頁範圍
        cover_end = 0
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            if txt in ('摘要', 'Abstract', '致謝', '目錄', '圖目錄', '表目錄'):
                cover_end = i
                break
        if cover_end == 0:
            return []  # 找不到封面範圍，跳過
        
        cover = doc.paragraphs[:cover_end]
        
        # 從 YAML 讀取封面關鍵字
        cover_kw = self.rule_engine.get('cover.detect_keywords', ['雷射', 'Study on', 'MAR'])
        
        # 找各區塊位置
        idx_top = []      # 校名、系所
        idx_thesis = []   # 碩士論文
        idx_title = []    # 中英文題目
        idx_bottom = []   # 研究生、指導教授、日期
        
        for i, p in enumerate(cover):
            txt = p.text.strip()
            if not txt:
                continue
            if '大學' in txt or '學系' in txt or '碩士班' in txt:
                idx_top.append(i)
            elif txt == '碩士論文':
                idx_thesis.append(i)
            elif '研究生' in txt or '指導教授' in txt or '中華民國' in txt:
                idx_bottom.append(i)
            elif any(kw in txt for kw in cover_kw):
                idx_title.append(i)
        
        def set_space_after(p, pt_val):
            pPr = p._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._element.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:after'), str(int(pt_val * 20)))  # twips
        
        def set_space_before(p, pt_val):
            pPr = p._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._element.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:before'), str(int(pt_val * 20)))
        
        def ensure_empty_para_after(doc, ref_para, cover_end_idx):
            """在 ref_para 後插入一個空段落（如果後面不是空段落）"""
            ref_idx = None
            for i, p in enumerate(doc.paragraphs):
                if p._element is ref_para._element:
                    ref_idx = i
                    break
            if ref_idx is None:
                return
            next_idx = ref_idx + 1
            if next_idx < len(doc.paragraphs) and doc.paragraphs[next_idx].text.strip() == '':
                return  # 已有空段落
            # 在 ref_para 後插入空段落
            new_p = OxmlElement('w:p')
            ref_para._element.addnext(new_p)
        
        # 區塊1：校名+系所 → 碩士論文之間加空格
        if idx_top and idx_thesis:
            last_top = max(idx_top)
            # 在最後一個 top 段落後加空段落
            ensure_empty_para_after(doc, cover[last_top], cover_end)
            count += 1
        
        # 區塊2：碩士論文 → 題目之間加空格
        if idx_thesis and idx_title:
            last_thesis = max(idx_thesis)
            ensure_empty_para_after(doc, cover[last_thesis], cover_end)
            count += 1
        
        # 區塊3：題目 → 底部之間加空格
        if idx_title and idx_bottom:
            last_title = max(idx_title)
            ensure_empty_para_after(doc, cover[last_title], cover_end)
            count += 1
        
        # 設定封面頁 section 垂直對齊為 both
        for section in doc.sections:
            sectPr = section._sectPr
            vAlign = sectPr.find(qn('w:vAlign'))
            if vAlign is None:
                vAlign = OxmlElement('w:vAlign')
                sectPr.append(vAlign)
            vAlign.set(qn('w:val'), 'both')
            count += 1
            break
        
        # 設定封面頁字型大小
        for i, p in enumerate(cover):
            txt = p.text.strip()
            if not txt:
                continue
            # 判斷是否為題目
            is_title = any(kw in txt for kw in cover_kw)
            
            for run in p.runs:
                rPr = run._element.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    run._element.insert(0, rPr)
                sz = rPr.find(qn('w:sz'))
                if sz is None:
                    sz = OxmlElement('w:sz')
                    rPr.append(sz)
                szCs = rPr.find(qn('w:szCs'))
                if szCs is None:
                    szCs = OxmlElement('w:szCs')
                    rPr.append(szCs)
                if is_title:
                    sz.set(qn('w:val'), '48')    # 24pt = 48 half-pt
                    szCs.set(qn('w:val'), '48')
                    # 題目粗體
                    b = rPr.find(qn('w:b'))
                    if b is None:
                        b = OxmlElement('w:b')
                        rPr.append(b)
                else:
                    sz.set(qn('w:val'), '48')    # 24pt = 48 half-pt
                    szCs.set(qn('w:val'), '48')
                    # 設定粗體
                    b = rPr.find(qn('w:b'))
                    if b is None:
                        b = OxmlElement('w:b')
                        rPr.append(b)
                # 設定字型
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), body_eng)
                rFonts.set(qn('w:hAnsi'), body_eng)
                rFonts.set(qn('w:eastAsia'), body_font)
            count += 1
        
        if count:
            return [{'type': 'cover_tabs', 'message': f'已調整封面頁{count}處區塊間距'}]
        return []
