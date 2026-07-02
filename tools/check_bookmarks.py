import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

# 檢查真實論文是否有 ref_ bookmarks
doc = Document(r'論文(最終版)-昱翔 - 複製.docx')
found = False
for p in doc.paragraphs:
    for bm in p._element.findall('.//' + qn('w:bookmarkStart')):
        name = bm.get(qn('w:name'), '')
        if name.startswith('ref_'):
            print(f'Found: {name}')
            found = True
if not found:
    print('No ref_ bookmarks found in original document')

print()

# 檢查修正後的
doc2 = Document(r'test_output/real_thesis_fixed_v5.docx')
found2 = False
for p in doc2.paragraphs:
    for bm in p._element.findall('.//' + qn('w:bookmarkStart')):
        name = bm.get(qn('w:name'), '')
        if name.startswith('ref_'):
            print(f'Fixed doc: Found {name}')
            found2 = True
if not found2:
    print('Fixed doc: No ref_ bookmarks found')
