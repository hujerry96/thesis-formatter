import sys; sys.path.insert(0, '.')
from docx import Document

src = r'論文(最終版)-昱翔 - 複製.docx'
doc = Document(src)

# 找 "目錄" 附近的段落
found = False
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if '目錄' in txt or '表目' in txt or '圖目' in txt:
        print(f'[{i:4d}] [{para.style.name:20s}] {repr(txt[:80])}')

print('\n--- 目錄區域段落（推測在 Heading "目錄" 之後）---')
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if '目錄' == txt:
        found = True
        for j in range(i, min(i+30, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            pt = p.text.strip()
            if not pt:
                continue
            if 'Heading' in (p.style.name or ''):
                break
            print(f'[{j:4d}] [{p.style.name:20s}] {repr(pt[:80])}')
        break
