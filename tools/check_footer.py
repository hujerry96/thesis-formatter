import sys; sys.path.insert(0, '.')
from docx import Document
from lxml import etree

qn = lambda x: '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + x

doc = Document(r'論文(最終版)-昱翔 - 複製.docx')

for i, section in enumerate(doc.sections):
    footer = section.footer
    ftr_elem = footer._element
    paragraphs = footer.paragraphs
    print(f'=== 第{i+1}節 ===')
    print(f'  is_linked_to_previous = {footer.is_linked_to_previous}')
    print(f'  len(footer.paragraphs) = {len(paragraphs)}')
    # 看 footer XML 結構
    xml_str = etree.tostring(ftr_elem, pretty_print=True, encoding='unicode')
    print(f'  footer XML (前300字): {xml_str[:300]}')

    # 檢查 sectPr 的 footerReference
    sectPr = section._sectPr
    for ref in sectPr.findall(qn('w:footerReference')):
        ref_type = ref.get(qn('w:type'), 'default')
        ref_id = ref.get(qn('r:id'), 'none')
        print(f'  w:footerReference type={ref_type}, id={ref_id}')

    print()
