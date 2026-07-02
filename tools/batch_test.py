"""批次測試腳本：測試 3 篇論文 + 驗證修正結果"""
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from core.analyzer import FormatAnalyzer
from core.fixer import FormatFixer

PROJECT_ROOT = Path(__file__).parent.parent
TEST_DIR = PROJECT_ROOT / "test_papers"
OUT_DIR = PROJECT_ROOT / "test_output"
OUT_DIR.mkdir(parents=True, exist_ok=True)

RULE_FILE = "thesis_zh.yaml"


def validate_document(doc_path: Path) -> list:
    """驗證修正後的文件是否有常見問題"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_LINE_SPACING

    doc = Document(str(doc_path))
    errors = []

    for i, section in enumerate(doc.sections):
        footer = section.footer
        for para in footer.paragraphs:
            page_count = 0
            for run in para._element.findall('.//' + qn('w:r')):
                fld = run.find(qn('w:fldChar'))
                instr = run.find(qn('w:instrText'))
                if fld is not None and fld.get(qn('w:fldCharType')) == 'begin':
                    page_count += 1
            if page_count > 1:
                errors.append(f'第{i+1}節頁尾發現 {page_count} 組 PAGE 欄位（應為 1 組）')

        page_code_count = len(para._element.findall('.//' + qn('w:instrText')))
        if page_code_count > 1:
            errors.append(f'第{i+1}節頁尾有 {page_code_count} 個 instrText（可能重複頁碼）')

    for i, para in enumerate(doc.paragraphs):
        if para.style and 'Heading 1' in para.style.name and para.text.strip():
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pb = pPr.find(qn('w:pageBreakBefore'))
                if pb is None or pb.get(qn('w:val')) == '0':
                    errors.append(f'Heading 1 "{para.text.strip()[:20]}" 缺少分頁符號')

        txt = para.text.strip()
        if 'Heading' in (para.style.name if para.style else ''):
            continue
        if not txt:
            continue
        pf = para.paragraph_format
        if pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
            ls = pf.line_spacing
            if isinstance(ls, (int, float)) and abs(ls - 1.5) < 0.1:
                continue
            errors.append(f'段落 {i+1} 行距為 {ls}，應為 1.5')

    return errors


def count_pages_approx(doc_path: Path) -> int:
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(doc_path))
    page_breaks = 0
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            pb = pPr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                page_breaks += 1
        for r in para._element.findall('.//' + qn('w:br')):
            if r.get(qn('w:type')) == 'page':
                page_breaks += 1
    sec_count = len(doc.sections) - 1
    return max(1, page_breaks + sec_count + 1)


def run_test(paper_path: Path, label: str):
    print(f'\n{"="*60}')
    print(f'測試: {label}')
    print(f'檔案: {paper_path.name}')
    print(f'{"="*60}')

    analyzer = FormatAnalyzer(RULE_FILE)
    fixer = FormatFixer(RULE_FILE)

    t0 = time.time()
    issues = analyzer.analyze(str(paper_path))
    t1 = time.time()
    print(f'分析完成: {len(issues)} 個問題 ({t1-t0:.2f}s)')

    out_path = OUT_DIR / f'{paper_path.stem}_修正{paper_path.suffix}'
    t0 = time.time()
    fixes = fixer.fix_document(str(paper_path), str(out_path))
    t1 = time.time()
    print(f'修正完成: {len(fixes)} 項修正 ({t1-t0:.2f}s)')

    print(f'輸出: {out_path.name}')

    sev_count = {'error': 0, 'warning': 0, 'info': 0}
    for issue in issues:
        s = issue.get('severity', 'info')
        if s in sev_count:
            sev_count[s] += 1
    print(f'問題統計: 錯誤={sev_count["error"]} 警告={sev_count["warning"]} 提示={sev_count["info"]}')

    validation_errors = validate_document(out_path)
    if validation_errors:
        print(f'\n[FAIL] 驗證發現 {len(validation_errors)} 個問題:')
        for ve in validation_errors:
            print(f'  - {ve}')
    else:
        print('\n[PASS] 文件驗證通過')

    orig_approx_pages = count_pages_approx(paper_path)
    fixed_approx_pages = count_pages_approx(out_path)
    print(f'原始約 {orig_approx_pages} 頁 → 修正後約 {fixed_approx_pages} 頁')

    return {
        'name': label,
        'issues_count': len(issues),
        'fixes_count': len(fixes),
        'validation_errors': len(validation_errors),
        'issues': issues,
        'fixes': fixes,
        'validation': validation_errors,
        'output': out_path,
    }


def main():
    papers = [
        (TEST_DIR / 'test_paper_A_general.docx', '論文 A - 一般中文論文'),
        (TEST_DIR / 'test_paper_B_complex.docx', '論文 B - 複雜引用論文'),
    ]
    edge = TEST_DIR / 'test_paper_C_edge.docx'
    if edge.exists():
        papers.append((edge, '論文 C - 邊界案例'))

    results = []
    for path, label in papers:
        if not path.exists():
            print(f'⚠ 跳過 {label}: {path.name} 不存在')
            continue
        result = run_test(path, label)
        results.append(result)

    print(f'\n{"="*60}')
    print('測試摘要')
    print(f'{"="*60}')
    for r in results:
        status = '[PASS]' if r['validation_errors'] == 0 else '[FAIL]'
        print(f'{status} {r["name"]}: {r["issues_count"]} 問題 / {r["fixes_count"]} 修正 / {r["validation_errors"]} 驗證錯誤')

    total_ve = sum(r['validation_errors'] for r in results)
    if total_ve == 0:
        print('\n全部測試通過！')
    else:
        print(f'\n共 {total_ve} 個驗證錯誤需要修正')


if __name__ == '__main__':
    main()
