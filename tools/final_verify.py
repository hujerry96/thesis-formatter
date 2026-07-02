import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET

doc = Document(r'test_output/real_thesis_fixed_v5.docx')

print(f'=== Section 數量: {len(doc.sections)} ===')

h1 = [(i, p.text[:60], p.style.name) for i, p in enumerate(doc.paragraphs) if 'Heading 1' in (p.style.name or '')]
print(f'Heading 1: {len(h1)}')
for i, t, s in h1: print(f'  [{i}] {t}')

h2 = [(p.text[:60], p.style.name) for p in doc.paragraphs if 'Heading 2' in (p.style.name or '')]
print(f'Heading 2: {len(h2)}')

print(f'Tables: {len(doc.tables)}')

for i, sec in enumerate(doc.sections):
    ftr = sec.footer._element
    pgnum = sec._sectPr.find(qn('w:pgNumType'))
    start = pgnum.get(qn('w:start')) if pgnum is not None else 'none'
    sdt_count = len(ftr.findall(qn('w:sdt')))
    xml = ET.tostring(ftr, encoding='unicode')
    import re
    instrs = re.findall(r'<w:instrText[^>]*>([^<]+)', xml)
    print(f'Section {i}: start={start}, SDT={sdt_count}, instr={instrs}')
