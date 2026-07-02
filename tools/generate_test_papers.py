"""生成 3 篇 60 頁工程領域測試論文（含多種格式問題）"""
import random
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

random.seed(42)

PROJECT_ROOT = Path(__file__).parent.parent
OUT_DIR = PROJECT_ROOT / "test_papers"
OUT_DIR.mkdir(parents=True, exist_ok=True)

ENG_CHAPTERS = [
    "Introduction", "Literature Review", "Research Methodology",
    "Experimental Results and Discussion", "Conclusion and Future Work"
]
ZH_CHAPTERS = [
    "緒論", "文獻回顧", "研究方法",
    "實驗結果與討論", "結論與未來展望"
]

ZH_PARAS = [
    "近年來隨著高科技產業的快速發展，材料科學與工程領域面臨前所未有的挑戰與機遇。",
    "本研究旨在探討高性能合金材料在極端環境下的微觀結構演化與機械性能變化。",
    "實驗結果顯示，經熱處理後的試樣其抗拉強度較原始狀態提升約百分之十五。",
    "掃描式電子顯微鏡觀察結果表明，析出相的尺寸與分佈對材料性能有顯著影響。",
    "本研究所提出的模型能夠有效預測材料在高溫下的潛變行為與疲勞壽命。",
    "透過 X 光繞射分析，可以確認不同熱處理條件下晶體結構的變化趨勢。",
    "與傳統製程相比，本研究所採用的新穎製程能夠大幅降低生產成本與能源消耗。",
    "有限元素分析結果與實驗數據吻合良好，驗證了數值模型的準確性與可靠性。",
    "表面改質技術的應用顯著提升了材料的耐磨耗性能與抗腐蝕能力。",
    "本研究之成果可應用於航太、能源及生醫等高端工程領域，具有重要的學術與實用價值。",
    "在過去數十年中，學者們對金屬材料的強化機制進行了廣泛而深入的研究。",
    "缺陷工程已被證明是調控材料性能的有效策略之一，受到學術界的高度重視。",
    "本實驗採用真空電弧熔煉技術製備高熵合金，確保成分均勻性與純度。",
    "熱力學計算與動力學模擬的結合為材料設計提供了強而有力的理論依據。",
    "機械性質測試包括硬度、拉伸、衝擊及疲勞等項目，涵蓋材料的各項關鍵性能指標。",
    "利用穿透式電子顯微鏡進行微觀結構分析，可觀察到奈米尺度的析出相與差排結構。",
    "晶粒細化是提升材料強度的重要手段之一，符合 Hall-Petch 關係式的預測。",
    "腐蝕行為評估採用電化學阻抗譜與極化曲線量測，全面分析材料的耐蝕性能。",
    "優化製程參數後，材料的綜合性能得到顯著改善，符合工業應用標準。",
    "本研究的創新點在於首次將機器學習技術應用於合金成分的最佳化設計。",
]

EN_PARAS = [
    "Advanced manufacturing technologies have revolutionized the production of high-performance materials.",
    "The mechanical properties of metallic alloys are strongly influenced by their microstructural characteristics.",
    "In situ characterization techniques provide valuable insights into the deformation mechanisms.",
    "Grain boundary engineering has emerged as a promising approach for enhancing material performance.",
    "The development of lightweight materials is crucial for improving fuel efficiency in aerospace applications.",
    "Computational modeling plays an increasingly important role in predicting material behavior.",
    "The correlation between processing parameters and resulting properties is a key area of investigation.",
    "Surface modification techniques offer effective solutions for improving wear and corrosion resistance.",
    "The integration of experimental and computational approaches accelerates the materials discovery process.",
    "Recent advances in additive manufacturing have opened new possibilities for complex component fabrication.",
    "The Hall-Petch relationship describes the strengthening effect of grain size reduction on yield strength.",
    "Phase transformations during heat treatment significantly affect the final mechanical properties.",
    "Nanoscale precipitates play a crucial role in enhancing both strength and ductility of structural alloys.",
    "Thermodynamic stability of the constituent phases determines the long-term performance under service conditions.",
    "Fatigue crack propagation behavior is strongly dependent on the microstructure and loading conditions.",
    "Electrochemical impedance spectroscopy provides a powerful tool for evaluating corrosion resistance.",
    "Solid solution strengthening is one of the fundamental mechanisms for improving alloy strength.",
    "Texture evolution during thermomechanical processing influences the anisotropy of mechanical properties.",
    "The application of high-entropy alloys has expanded the compositional space for material design.",
    "Fracture toughness measurements provide critical data for structural integrity assessments.",
]


def set_cell_text(cell, text, bold=False, font_name='新細明體', size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_field_page_number(para, style='ARABIC'):
    """在段落中插入 PAGE 功能變數"""
    fld_begin = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    fld_begin.append(fc1)
    para._element.append(fld_begin)

    fld_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' PAGE \\* {style} '
    fld_instr.append(instr)
    para._element.append(fld_instr)

    fld_sep = OxmlElement('w:r')
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    fld_sep.append(fc_sep)
    para._element.append(fld_sep)

    fld_end = OxmlElement('w:r')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    fld_end.append(fc_end)
    para._element.append(fld_end)


def add_footer_page_number(section, style='ARABIC', start_from=None):
    """為節加入頁尾頁碼"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field_page_number(p, style)
    if start_from is not None:
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        pgNumType.set(qn('w:start'), str(start_from))


def add_toc_field(doc):
    """插入目錄 TOC 功能變數"""
    p = doc.add_paragraph()
    run = p.add_run('目錄')
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.name = '標楷體'
    run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fld_begin = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    fld_begin.append(fc1)
    p2._element.append(fld_begin)

    fld_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_instr.append(instr)
    p2._element.append(fld_instr)

    fld_sep = OxmlElement('w:r')
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    fld_sep.append(fc_sep)
    p2._element.append(fld_sep)

    fld_end = OxmlElement('w:r')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    fld_end.append(fc_end)
    p2._element.append(fld_end)


def add_heading_para(doc, text, level=1, buggy=False):
    """加入標題段落，buggy=True 時故意不加分頁"""
    p = doc.add_paragraph(text)
    if level == 1:
        p.style = doc.styles['Heading 1']
    elif level == 2:
        p.style = doc.styles['Heading 2']
    else:
        p.style = doc.styles['Heading 3']

    if not buggy:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p._element.insert(0, pPr)
        pb = pPr.find(qn('w:pageBreakBefore'))
        if pb is None:
            pb = OxmlElement('w:pageBreakBefore')
            pPr.append(pb)
    return p


def add_body_para(doc, text, buggy_spacing=False, buggy_font=False):
    """加入內文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)

    if buggy_font:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    else:
        run.font.name = '新細明體'
        run.font.size = Pt(12)

    if buggy_spacing:
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5

    pf.first_line_indent = Cm(0.75)

    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '新細明體')
    rFonts.set(qn('w:ascii'), 'Times New Roman')

    return p


def add_buggy_caption(doc, prefix, num, text):
    """加入 buggy 格式的圖表說明（使用 - 而非 .）"""
    p = doc.add_paragraph()
    run = p.add_run(f'{prefix} {num} {text}')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_ref_table(doc, refs):
    """加入參考文獻表格（常見格式問題來源）"""
    p = doc.add_paragraph('參考文獻')
    p.style = doc.styles['Heading 1']

    table = doc.add_table(rows=len(refs), cols=2)
    table.style = 'Table Grid'
    for i, (num, content) in enumerate(refs):
        set_cell_text(table.rows[i].cells[0], num)
        set_cell_text(table.rows[i].cells[1], content)


def generate_paper_a():
    """論文 A: 一般中文論文 — 封面+摘要+目錄+5章+圖表+參考文獻"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(3.2)

    add_footer_page_number(section, 'ARABIC')

    paras_needed = 0
    for zh, en in zip(ZH_CHAPTERS, ENG_CHAPTERS):
        paras_needed += max(35, (60 * len(ZH_PARAS) + len(EN_PARAS)) // len(ZH_CHAPTERS))
    total_paras = max(paras_needed, 160)

    p = doc.add_paragraph('國立清華大學')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('材料科學工程學系碩士班')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('')
    p = doc.add_paragraph('碩士論文')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('')
    p = doc.add_paragraph('雷射加工參數對高熵合金微觀結構與機械性質之研究')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('')
    p = doc.add_paragraph('研究生：張大中')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('')
    p = doc.add_paragraph('指導教授：王志明')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('')
    p = doc.add_paragraph('中華民國 114 年 7 月')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    p = doc.add_paragraph('摘要')
    p.style = doc.styles['Heading 1']
    for _ in range(8):
        txt = random.choice(ZH_PARAS) + random.choice(ZH_PARAS) + random.choice(ZH_PARAS)
        add_body_para(doc, txt)

    p = doc.add_paragraph('Abstract')
    p.style = doc.styles['Heading 1']
    for _ in range(6):
        txt = random.choice(EN_PARAS) + random.choice(EN_PARAS)
        add_body_para(doc, txt)

    doc.add_page_break()
    p = doc.add_paragraph('誌謝')
    p.style = doc.styles['Heading 1']
    for _ in range(3):
        add_body_para(doc, random.choice(ZH_PARAS) + random.choice(ZH_PARAS))

    doc.add_page_break()
    add_toc_field(doc)
    doc.add_page_break()

    para_idx = 0
    for ci, (zh_ch, en_ch) in enumerate(zip(ZH_CHAPTERS, ENG_CHAPTERS)):
        add_heading_para(doc, f'{ci+1}.  {zh_ch}', level=1, buggy=ci == 0)

        if ci == 0:
            add_heading_para(doc, f'{ci+1}.1  研究背景', level=2, buggy=True)
            for _ in range(4):
                add_body_para(doc, random.choice(ZH_PARAS) + random.choice(ZH_PARAS),
                              buggy_font=True)
            add_heading_para(doc, f'{ci+1}.2  研究動機與目的', level=2)
            for _ in range(4):
                add_body_para(doc, random.choice(ZH_PARAS) + random.choice(ZH_PARAS))

        if ci == 2:
            add_heading_para(doc, f'{ci+1}.1  實驗材料', level=2)
            for _ in range(3):
                add_body_para(doc, random.choice(ZH_PARAS) + random.choice(EN_PARAS))

            add_heading_para(doc, f'{ci+1}.2  實驗設備', level=2)
            add_buggy_caption(doc, '圖', '2-1', '掃描式電子顯微鏡與能譜分析儀')
            add_buggy_caption(doc, '圖', '2-2', 'X光繞射分析儀')
            add_body_para(doc, random.choice(ZH_PARAS))

            add_heading_para(doc, f'{ci+1}.3  實驗步驟', level=2)
            for pi in range(5):
                p = doc.add_paragraph(f'{pi+1}. {random.choice(ZH_PARAS)}')
                p.paragraph_format.first_line_indent = Cm(0)

            add_buggy_caption(doc, '表', '3-1', '實驗參數設定與對照組')
            for _ in range(2):
                add_body_para(doc, random.choice(ZH_PARAS))

            table = doc.add_table(rows=6, cols=4)
            table.style = 'Table Grid'
            headers = ['試樣編號', '溫度 (°C)', '時間 (hr)', '硬度 (HRC)']
            for ci_h, h in enumerate(headers):
                set_cell_text(table.rows[0].cells[ci_h], h, bold=True)
            for ri in range(1, 6):
                set_cell_text(table.rows[ri].cells[0], f'S-{ri:03d}')
                set_cell_text(table.rows[ri].cells[1], str(random.randint(800, 1200)))
                set_cell_text(table.rows[ri].cells[2], str(random.randint(1, 24)))
                set_cell_text(table.rows[ri].cells[3], f'{random.uniform(35, 55):.1f}')

        if ci == 3:
            add_heading_para(doc, f'{ci+1}.1  微觀結構分析', level=2)
            add_buggy_caption(doc, '圖', '4-1', '不同熱處理條件下之 SEM 影像')
            add_buggy_caption(doc, '圖', '4-2', 'EDS 元素分布圖')
            for _ in range(6):
                add_body_para(doc, random.choice(ZH_PARAS) + random.choice(EN_PARAS),
                              buggy_spacing=random.random() < 0.3)

            add_heading_para(doc, f'{ci+1}.2  機械性質測試', level=2)
            add_heading_para(doc, f'{ci+1}.2.1  硬度測試', level=3)
            for _ in range(3):
                add_body_para(doc, random.choice(ZH_PARAS))
            add_heading_para(doc, f'{ci+1}.2.2  拉伸測試', level=3)
            for _ in range(3):
                add_body_para(doc, random.choice(ZH_PARAS) + random.choice(EN_PARAS))

            add_buggy_caption(doc, '表', '4-1', '拉伸測試結果彙整')
            table = doc.add_table(rows=5, cols=4)
            table.style = 'Table Grid'
            headers = ['條件', '降伏強度 (MPa)', '抗拉強度 (MPa)', '伸長率 (%)']
            for ci_h, h in enumerate(headers):
                set_cell_text(table.rows[0].cells[ci_h], h, bold=True)
            for ri in range(1, 5):
                set_cell_text(table.rows[ri].cells[0], f'條件 {ri}')
                set_cell_text(table.rows[ri].cells[1], str(random.randint(300, 800)))
                set_cell_text(table.rows[ri].cells[2], str(random.randint(500, 1200)))
                set_cell_text(table.rows[ri].cells[3], f'{random.uniform(5, 30):.1f}')

        for _ in range(max(8, total_paras // len(ZH_CHAPTERS))):
            zh = random.choice(ZH_PARAS)
            en = random.choice(EN_PARAS) if random.random() < 0.4 else ''
            txt = zh + (' ' + en if en else '')
            buggy_s = random.random() < 0.15
            buggy_f = random.random() < 0.1
            p = add_body_para(doc, txt, buggy_spacing=buggy_s, buggy_font=buggy_f)
            para_idx += 1

    refs = [(f'[{i}]', f'Author {chr(64+i)}. ({2020 + i % 5}). Study on advanced materials processing. Journal of Materials Science, {random.randint(50, 60)}({random.randint(1, 12)}), {random.randint(100, 999)}-{random.randint(1000, 9999)}.') for i in range(1, 21)]
    add_ref_table(doc, refs)

    path = OUT_DIR / 'test_paper_A_general.docx'
    doc.save(str(path))
    print(f'論文 A 已產生: {path}')
    return path


def generate_paper_b():
    """論文 B: 複雜引用論文 — 多層 Heading、交叉引用、範圍引用"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    add_footer_page_number(section, 'ARABIC')

    for _ in range(10):
        doc.add_paragraph('')

    p = doc.add_paragraph('先進陶瓷複合材料之製程開發與性能評估')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(24)
    p.runs[0].font.bold = True

    p = doc.add_paragraph('國立臺灣科技大學')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('機械工程系')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('碩士論文')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('研究生：李建華')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('指導教授：陳茂林')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('中華民國 114 年 6 月')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    p = doc.add_paragraph('摘要')
    p.style = doc.styles['Title']
    for _ in range(6):
        add_body_para(doc, random.choice(ZH_PARAS) + '相關研究[1-3]已證實此現象。' + random.choice(ZH_PARAS))

    doc.add_page_break()
    p = doc.add_paragraph('目錄')
    p.style = doc.styles['Heading 1']
    add_toc_field(doc)
    doc.add_page_break()

    for ci in range(7):
        add_heading_para(doc, f'第{ci+1}章  {random.choice(ZH_CHAPTERS)}', level=1, buggy=ci == 3)

        add_heading_para(doc, f'{ci+1}.1 {random.choice(ZH_PARAS)[:10]}', level=2)
        for _ in range(5):
            ref_nums = random.sample(range(1, 26), random.randint(1, 3))
            ref_str = ','.join(str(n) for n in sorted(ref_nums))
            txt = random.choice(ZH_PARAS) + f'相關文獻[{ref_str}]' + random.choice(EN_PARAS)
            add_body_para(doc, txt)

        if random.random() < 0.6:
            add_heading_para(doc, f'{ci+1}.1.{random.randint(1,3)} {random.choice(ZH_PARAS)[:8]}', level=3)
            for _ in range(3):
                add_body_para(doc, random.choice(ZH_PARAS) + f'如文獻[11-15]所指出，' + random.choice(ZH_PARAS))

        if ci % 2 == 0:
            add_heading_para(doc, f'{ci+1}.2 {random.choice(ZH_PARAS)[:10]}', level=2)
            add_buggy_caption(doc, '圖', f'{ci+1}-1', f'陶瓷複合材料微觀結構 SEM 影像')
            for _ in range(4):
                txt = random.choice(ZH_PARAS) + f'如圖{ci+1}-1所示，' + random.choice(ZH_PARAS)
                add_body_para(doc, txt)

    refs = [(f'[{i}]', f'Author {chr(64+i)}. Study on ceramic matrix composites. Journal of Advanced Materials, 202{random.randint(0,4)}.') for i in range(1, 26)]
    add_ref_table(doc, refs)

    path = OUT_DIR / 'test_paper_B_complex.docx'
    doc.save(str(path))
    print(f'論文 B 已產生: {path}')
    return path


def generate_paper_c():
    """論文 C: 邊界案例 — 無封面、多節多頁碼、內文大量空段落"""
    doc = Document()

    section0 = doc.sections[0]
    add_footer_page_number(section0, 'roman', start_from=1)

    for _ in range(3):
        doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph('附錄 A：實驗數據補充資料')
    p.style = doc.styles['Heading 1']
    for _ in range(4):
        add_body_para(doc, random.choice(ZH_PARAS) + random.choice(EN_PARAS))

    p = doc.add_paragraph('A.1 XRD 原始數據')
    p.style = doc.styles['Heading 2']
    for _ in range(3):
        add_body_para(doc, random.choice(ZH_PARAS))

    table = doc.add_table(rows=15, cols=5)
    table.style = 'Table Grid'
    for ri, row in enumerate(table.rows):
        for ci in range(5):
            set_cell_text(row.cells[ci], str(random.randint(100, 999)) if ri > 0 else f'Hdr{ci+1}',
                          bold=(ri == 0))

    for _ in range(2):
        doc.add_paragraph('')

    add_buggy_caption(doc, '圖', 'A-1', 'XRD 繞射圖譜')

    doc.add_section()
    section1 = doc.sections[1]
    add_footer_page_number(section1, 'ARABIC', start_from=1)

    p = doc.add_paragraph('附錄 B：模擬程式碼')
    p.style = doc.styles['Heading 1']
    code_lines = [
        "def calculate_stress(strain, modulus):",
        "    stress = strain * modulus",
        "    return stress",
        "",
        "def main():",
        "    strains = [0.001, 0.002, 0.005, 0.01]",
        "    modulus = 200e9  # Pa",
        "    for s in strains:",
        "        sigma = calculate_stress(s, modulus)",
        "        print(f'Strain: {s}, Stress: {sigma:.2e} Pa')",
    ]
    for line in code_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_page_break()
    add_buggy_caption(doc, '表', 'B-1', '模擬輸入參數')

    doc.add_section()
    section2 = doc.sections[2]
    add_footer_page_number(section2, 'ARABIC')

    p = doc.add_paragraph('附錄 C：實驗儀器型號')
    p.style = doc.styles['Heading 1']

    for ci in range(3):
        add_heading_para(doc, f'C.{ci+1} {random.choice(ZH_PARAS)[:10]}', level=2)
        for _ in range(5):
            add_body_para(doc, random.choice(ZH_PARAS) + random.choice(EN_PARAS),
                          buggy_spacing=random.random() < 0.5,
                          buggy_font=random.random() < 0.3)

    for _ in range(15):
        doc.add_paragraph('')

    path = OUT_DIR / 'test_paper_C_edge.docx'
    doc.save(str(path))
    print(f'論文 C 已產生: {path}')
    return path


if __name__ == '__main__':
    p1 = generate_paper_a()
    p2 = generate_paper_b()
    p3 = generate_paper_c()
    print(f'\n已產生 3 篇測試論文:')
    for p in [p1, p2, p3]:
        sz = p.stat().st_size
        print(f'  {p.name}  ({sz/1024:.0f} KB)')
