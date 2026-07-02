import sys; sys.path.insert(0, '.')
from docx import Document
from core.fixer import FormatFixer
import re

qn = lambda x: '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + x
KW_RE = re.compile(r'^(關鍵詞|關鍵字|Key words|Keywords)\s*[:：]?\s*', re.IGNORECASE)

src = r'論文(最終版)-昱翔 - 複製.docx'

# 先確認原文有關鍵詞
doc = Document(src)
found = False
for i, para in enumerate(doc.paragraphs):
    if KW_RE.match(para.text.strip()):
        txt = para.text.strip()[:80]
        pf = para.paragraph_format
        indent = pf.first_line_indent
        align = pf.alignment
        print(f'找到關鍵詞: [{i}] {repr(txt)}')
        print(f'  首行縮排={indent}, 對齊={align}')
        found = True
        break
if not found:
    print('未找到關鍵詞')

# 執行修正
doc2 = Document(src)
fixer = FormatFixer()
result = fixer._fix_keywords_format(doc2)
print(f'\n修正結果: {result}')

# 驗證修正後
for para in doc2.paragraphs:
    if KW_RE.match(para.text.strip()):
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            jc = pPr.find(qn('w:jc'))
            sp = pPr.find(qn('w:spacing'))
            print(f'修正後 ind={ind.get(qn("w:firstLine")) if ind is not None else "無"}')
            print(f'修正後 jc={jc.get(qn("w:val")) if jc is not None else "無"}')
            print(f'修正後 before={sp.get(qn("w:before")) if sp is not None else "無"}')
        break
