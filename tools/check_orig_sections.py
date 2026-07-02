from docx import Document
from docx.oxml.ns import qn

doc = Document(r'論文(最終版)-昱翔 - 複製.docx')
print(f'Total sections: {len(doc.sections)}')
for i, sec in enumerate(doc.sections):
    print(f'\n--- Section {i} ---')
    sp = sec._sectPr
    pgnum = sp.find(qn('w:pgNumType'))
    if pgnum is not None:
        print(f'  pgNumType: start={pgnum.get(qn("w:start"))}')
    footer = sec.footer
    print(f'  footer linked_to_previous={footer.is_linked_to_previous}')
    for j, p in enumerate(footer.paragraphs):
        runs = p._element.findall(qn('w:r'))
        flds = [r for r in runs if r.find(qn('w:fldChar')) is not None or r.find(qn('w:instrText')) is not None]
        instr_texts = []
        for r in flds:
            it = r.find(qn('w:instrText'))
            if it is not None:
                instr_texts.append(it.text)
        print(f'  footer para {j}: {len(flds)} fld-related, instr_texts={instr_texts}')

# Check heading paragraphs to see section boundaries
print('\n=== Heading 1 boundaries ===')
for i, p in enumerate(doc.paragraphs):
    if 'Heading 1' in (p.style.name or ''):
        print(f'  para {i}: {p.text[:60]}')
