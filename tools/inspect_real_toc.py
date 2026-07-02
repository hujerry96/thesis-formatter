import sys; sys.path.insert(0, '.')
from docx import Document
import re

src = r'論文(最終版)-昱翔 - 複製.docx'
doc = Document(src)

print("=== 原始檔案中 toc 樣式段落 ===")
for i, para in enumerate(doc.paragraphs):
    sname = para.style.name if para.style else ''
    if 'toc' in sname.lower():
        txt = para.text.strip()[:100]
        print(f'[{i:4d}] [{sname}] {repr(txt[:80])}')
        if '-' in txt:
            print(f'        ^^^^ 包含連字號')

print("\n=== Heading 樣式段落 ===")
for i, para in enumerate(doc.paragraphs[:50]):
    sname = para.style.name if para.style else ''
    if 'heading' in sname.lower():
        txt = para.text.strip()[:100]
        print(f'[{i:4d}] [{sname}] {repr(txt[:80])}')
        if '-' in txt:
            print(f'        ^^^^ 包含連字號')
