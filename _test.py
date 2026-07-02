"""測試：範圍引用 [5-7]、[33.35] 轉 REF field，參考文獻靠右對齊"""
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn

doc = Document()

# 內文含各種引用格式
doc.add_paragraph('單一引用[1]和另一個[5]。')
doc.add_paragraph('範圍引用[5-7]，應該要有三個 REF field。')
doc.add_paragraph('逗號分隔[1,5]和[1,5,7]。')
doc.add_paragraph('句點格式[33.35]應該自動轉為[33-35]並產生 REF field。')

# 參考文獻
doc.add_paragraph('參考文獻').style = doc.styles['Heading 1']

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(13)

refs = [
    ('[1]', 'Author A. (2024). Paper one. Journal, 10(1), 1-10.'),
    ('[5]', 'Author E. (2023). Paper five. Journal, 10(5), 50-59.'),
    ('[6]', 'Author F. (2023). Paper six. Journal, 10(6), 60-69.'),
    ('[7]', 'Author G. (2023). Paper seven. Journal, 10(7), 70-79.'),
    ('[33]', 'Author CC. (2020). Paper thirty-three. Journal, 20(3), 330-339.'),
    ('[34]', 'Author DD. (2020). Paper thirty-four. Journal, 20(4), 340-349.'),
    ('[35]', 'Author EE. (2020). Paper thirty-five. Journal, 20(5), 350-359.'),
]
for i, (num, content) in enumerate(refs):
    table.rows[i].cells[0].text = num
    table.rows[i].cells[1].text = content

doc.save('test_range.docx')
print('test_range.docx created')
